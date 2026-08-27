from __future__ import annotations

import io
import re
import zipfile
from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from research_platform.figure_analysis import FigureObservation, GeneratedResearchFigure
from research_platform.report_synthesis import (
    StudyProfile,
    SynthesisPackage,
    SynthesisSection,
)
from research_platform.word_report import _figure_matches_section, build_word_report


def _body_text_in_order(document: Document) -> str:
    return "\n".join(
        node.text or "" for node in document.element.body.iter(qn("w:t"))
    )


def test_figure_section_matching_does_not_duplicate_a_comparison_chart() -> None:
    assert _figure_matches_section(
        "Bulgular ve karşılaştırmalı sonuçlar",
        "Bulgular ve karşılaştırmalı sonuçlar",
    )
    assert not _figure_matches_section(
        "Bulgular ve karşılaştırmalı sonuçlar",
        "Temel bulgular",
    )


def test_word_report_is_a_sourced_docx_with_embedded_figures() -> None:
    source = SimpleNamespace(
        id="source-1",
        title="Independent clinical source",
        family="academic",
        url="https://example.org/clinical",
    )
    claim = SimpleNamespace(
        id="claim-1",
        text="The method improved the measured outcome.",
        status="supported",
        audit={"question_relevance": 0.91},
    )
    evidence = SimpleNamespace(quote="The measured outcome improved by ten percent.")
    synthesis_package = SynthesisPackage(
        executive_summary=(
            "Across the available study, the measured outcome improved, although replication "
            "is still required [S01]."
        ),
        sections=[
            SynthesisSection(
                title="Measured outcome",
                synthesis="The available evidence reports an improved measured outcome [S01].",
                implications="The finding warrants independent replication [S01].",
                source_ids=["S01"],
                claim_ids=["claim-1"],
            )
        ],
        cross_study_assessment="Only one study context is available [S01].",
        conclusion="The result is promising but not yet independently replicated [S01].",
        uncertainty="External validation remains necessary [S01].",
        study_profiles=[
            StudyProfile(
                source_id="source-1",
                source_label="S01",
                title="Independent clinical source",
                contribution="External validation",
                evidence_design="External validation",
            )
        ],
        generated_by_llm=True,
    )
    figure_stream = io.BytesIO()
    Image.new("RGB", (900, 500), "white").save(figure_stream, format="PNG")
    figure_observation = FigureObservation(
        source_id="source-1",
        source_version_id="version-1",
        source_label="S01",
        source_title="Independent clinical source",
        image_hash="a" * 64,
        image_key="runs/test/figure.png",
        page_number=3,
        caption="Figure 1. Outcome.",
        figure_type="bar chart",
        title="Measured outcome",
        axes={"x": "Group", "y": "Outcome"},
        series=["Outcome"],
        data_points=[],
        flow_steps=[],
        main_findings=[],
        limitations=["Only one source figure is available."],
        recommended_section="Measured outcome",
        relevance_score=0.9,
        exact_values_visible=False,
        confidence=0.8,
        vision_model="qwen3.5:4b",
        include_in_report=True,
        selection_reason="The source figure visualises the measured outcome.",
    )
    generated_figure = GeneratedResearchFigure(
        name="17a_source_figure_excerpt.png",
        data=figure_stream.getvalue(),
        title="Measured outcome",
        caption="Source figure: measured outcome [S01, p. 3].",
        description="Original source figure excerpt.",
        section_title="Measured outcome",
        source_labels=["S01"],
        observation_hash="a" * 64,
        origin="source_excerpt",
        attribution="Independent clinical source — https://example.org/clinical",
        rights_statement="Internal research review; verify rights before distribution.",
    )

    report = build_word_report(
        run_id="01WORDTEST",
        title="Word exporter acceptance",
        question="Does the method improve the measured outcome?",
        language="en",
        coverage={
            "source_family_coverage": 1.0,
            "query_branch_coverage": 0.9,
            "claim_audit_coverage": 1.0,
            "estimated_completeness": 0.8,
            "unresolved_major_claims": 0,
        },
        sources=[source],
        claims=[claim],
        reportable_claims=[claim],
        evidence_by_claim={claim.id: [(evidence, source)]},
        executive_summary="One audited claim is supported by an independent source.",
        narrative="The result is reported only within the available evidence.",
        uncertainty="Replication and external validation remain necessary.",
        synthesis_package=synthesis_package,
        figure_observations=[figure_observation],
        research_figures=[generated_figure],
    )

    assert set(report.figures) == {
        "16a_research_contribution_landscape.png",
        "16b_theme_evidence_map.png",
    }
    with zipfile.ZipFile(io.BytesIO(report.document)) as archive:
        assert "word/document.xml" in archive.namelist()
        assert len([name for name in archive.namelist() if name.startswith("word/media/")]) == 3
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "Research contribution types" in document_xml
        assert "Theme-evidence map" in document_xml

    document = Document(io.BytesIO(report.document))
    level_one_headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
    ]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Word exporter acceptance" in text
    assert "Thematic evidence synthesis" in text
    assert "platform-performance charts" in text
    assert level_one_headings == [
        "Contents",
        "1. Summary",
        "2. Research frame",
        "3. Thematic evidence synthesis",
        "4. Cross-study assessment and conclusion",
        "Appendix A. Method, coverage, and reproducibility",
        "Appendix B. Thematic Literature Landscape",
        "Appendix C. Complete source catalog",
        "Appendix D. Audited claim register",
        "Appendix E. Source figure observation register",
    ]
    ordered_text = _body_text_in_order(document)
    topic_map_start = ordered_text.index("Appendix B. Thematic Literature Landscape")
    source_catalog_start = ordered_text.index("Appendix C. Complete source catalog")
    topic_map_text = ordered_text[topic_map_start:source_catalog_start]
    assert "Figure B.1. Literature landscape by study purpose." in topic_map_text
    assert "Figure B.2. A blue cell indicates" in topic_map_text
    assert "Research contribution" in topic_map_text
    assert "3. Thematic literature landscape" not in text
    assert "Appendix E. Source figure observation register" in text
    assert "Model interpretation" in text
    assert "verify rights before distribution" in text
    assert len(document.tables) >= 3


def _anchors_and_bookmarks(document_bytes: bytes) -> tuple[set[str], set[str]]:
    """Extract internal link targets and bookmark names from a .docx payload."""
    with zipfile.ZipFile(io.BytesIO(document_bytes)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    anchors = set(re.findall(r'w:anchor="([^"]+)"', xml))
    bookmarks = set(re.findall(r'w:bookmarkStart[^>]*w:name="([^"]+)"', xml))
    return anchors, bookmarks


def _minimal_report_inputs() -> dict:
    source = SimpleNamespace(
        id="source-1",
        title="Independent clinical source",
        family="academic",
        url="https://example.org/clinical",
    )
    claim = SimpleNamespace(
        id="claim-1",
        text="The method improved the measured outcome.",
        status="supported",
        audit={"question_relevance": 0.91},
    )
    evidence = SimpleNamespace(quote="The measured outcome improved by ten percent.")
    return {
        "run_id": "01LINKTEST",
        "title": "Source cross-reference",
        "question": "Does the method improve the measured outcome?",
        "language": "en",
        "coverage": {
            "source_family_coverage": 1.0,
            "query_branch_coverage": 0.9,
            "claim_audit_coverage": 1.0,
            "estimated_completeness": 0.8,
            "unresolved_major_claims": 0,
        },
        "sources": [source],
        "claims": [claim],
        "reportable_claims": [claim],
        "evidence_by_claim": {claim.id: [(evidence, source)]},
        "executive_summary": "One audited claim is supported by an independent source.",
        "narrative": "The result is reported only within the available evidence.",
        "uncertainty": "Replication and external validation remain necessary.",
    }


def test_unlocalized_figure_interpretation_is_hidden_and_audit_uses_dashes() -> None:
    inputs = _minimal_report_inputs()
    inputs.update(
        {
            "title": "Figür dili testi",
            "question": "Model performansı nasıl karşılaşıyor?",
            "language": "tr",
            "executive_summary": "Rapor dili Türkçedir.",
            "narrative": "Figür metinleri rapor diliyle aynı olmalıdır.",
            "uncertainty": "Ek doğrulama gereklidir.",
        }
    )
    image_stream = io.BytesIO()
    Image.new("RGB", (900, 500), "white").save(image_stream, format="PNG")
    observation = FigureObservation(
        source_id="source-1",
        source_version_id="version-1",
        source_label="S01",
        source_title="Independent clinical source",
        image_hash="z" * 64,
        image_key="runs/test/figure.png",
        page_number=3,
        caption="Figure 2. Reader performance.",
        figure_type="scatter_plot",
        title="Şekil 2. Okuyucu performansı",
        axes={"x": "Readers", "y": "Accuracy"},
        series=["Model"],
        data_points=[],
        flow_steps=[],
        main_findings=[],
        limitations=[],
        recommended_section="Karşılaştırmalı sonuçlar",
        relevance_score=0.9,
        exact_values_visible=False,
        confidence=0.9,
        vision_model="qwen3.5:4b#figure-v5",
        include_in_report=True,
        selection_reason="",
    )
    generated = GeneratedResearchFigure(
        name="17a_source_figure_excerpt.png",
        data=image_stream.getvalue(),
        title=observation.title,
        caption="Kaynak figürü: Şekil 2. Ayrıntılı özgün açıklama kaynak kaydında korunmuştur. [S01, s. 3].",
        description="S01 kaynağından kırpılan özgün araştırma figürü.",
        section_title="Karşılaştırmalı sonuçlar",
        source_labels=["S01"],
        observation_hash=observation.image_hash,
        origin="source_excerpt",
    )

    report = build_word_report(
        **inputs,
        figure_observations=[observation],
        research_figures=[generated],
    )
    document = Document(io.BytesIO(report.document))
    table_texts = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    full_text += "\n" + "\n".join(table_texts)

    assert "Model yorumu" not in full_text
    assert "The model outperformed" not in full_text
    assert sum(cell.count("—") for cell in table_texts) >= 2


def test_synthesis_report_links_citations_to_the_source_catalog() -> None:
    """A reader clicking [S01] in the prose should land on that row of the catalog."""
    package = SynthesisPackage(
        executive_summary="The outcome improved [S01].",
        sections=[
            SynthesisSection(
                title="Measured outcome",
                synthesis="The measured outcome improved in the available study [S01].",
                consensus="One study reports improvement [S01].",
                source_ids=["S01"],
                claim_ids=["claim-1"],
            )
        ],
        study_profiles=[
            StudyProfile(
                source_id="source-1",
                source_label="S01",
                title="Independent clinical source",
                contribution="Detection / diagnosis",
                evidence_design="Observational",
            )
        ],
        cross_study_assessment="Only one study is available [S01].",
        conclusion="Replication remains necessary [S01].",
        uncertainty="Single-study evidence [S01].",
        generated_by_llm=True,
    )
    report = build_word_report(**_minimal_report_inputs(), synthesis_package=package)
    anchors, bookmarks = _anchors_and_bookmarks(report.document)

    assert "src_S01" in bookmarks, "catalog row must be bookmarked"
    assert "src_S01" in anchors, "citations must link to it"


def test_turkish_synthesis_report_uses_ozet_heading() -> None:
    package = SynthesisPackage(
        executive_summary="Ölçülen sonuç iyileşmiştir [S01].",
        sections=[
            SynthesisSection(
                title="Ölçülen sonuç",
                synthesis="Mevcut kanıt iyileşme bildirmektedir [S01].",
                source_ids=["S01"],
                claim_ids=["claim-1"],
            )
        ],
        study_profiles=[],
        cross_study_assessment="Tek çalışma bağlamı bulunmaktadır [S01].",
        conclusion="Bağımsız doğrulama gereklidir [S01].",
        uncertainty="Kanıt tek çalışmayla sınırlıdır [S01].",
        generated_by_llm=True,
    )
    inputs = _minimal_report_inputs()
    inputs["language"] = "tr"

    report = build_word_report(**inputs, synthesis_package=package)
    document = Document(io.BytesIO(report.document))
    level_one_headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
    ]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "1. Özet" in text
    assert "Yönetici sentezi" not in text
    assert level_one_headings == [
        "İçindekiler",
        "1. Özet",
        "2. Araştırma çerçevesi",
        "3. Tematik kanıt sentezi",
        "4. Çalışmalar arası değerlendirme ve sonuç",
        "Ek A. Yöntem, kapsam ve yeniden üretilebilirlik",
        "Ek B. Literatürün Konu Haritası",
        "Ek C. Tam kaynak kataloğu",
        "Ek D. Denetlenmiş iddia kaydı",
    ]
    assert text.count("Ek B. Literatürün Konu Haritası") == 1
    assert "Şekil B.1. Çalışmaların araştırma amacına göre literatür görünümü." in text
    assert "Şekil B.2. Mavi hücre" in text
    assert "3. Literatürün konu haritası" not in text
    assert "Ek E. Kaynak figürü inceleme kaydı" not in text


def test_english_synthesis_report_uses_summary_heading() -> None:
    package = SynthesisPackage(
        executive_summary="The measured outcome improved [S01].",
        sections=[],
        study_profiles=[],
        cross_study_assessment="One study is available [S01].",
        conclusion="Independent validation is required [S01].",
        uncertainty="Evidence is limited to one study [S01].",
        generated_by_llm=True,
    )

    report = build_word_report(**_minimal_report_inputs(), synthesis_package=package)
    document = Document(io.BytesIO(report.document))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "1. Summary" in text
    assert "Executive synthesis" not in text


def test_focused_answer_report_also_bookmarks_its_catalog() -> None:
    """The two templates build separate catalogs; fixing one does not fix the other."""
    report = build_word_report(**_minimal_report_inputs(), synthesis_package=None)
    _, bookmarks = _anchors_and_bookmarks(report.document)
    assert "src_S01" in bookmarks


def test_report_has_no_links_pointing_at_missing_bookmarks() -> None:
    """A link to an undefined bookmark is silently inert in Word — worse than plain text."""
    package = SynthesisPackage(
        executive_summary="Two labels appear but only one source exists [S01] [S07].",
        sections=[],
        study_profiles=[],
        cross_study_assessment="",
        conclusion="Unknown label [S07] must stay plain text.",
        uncertainty="",
        generated_by_llm=True,
    )
    report = build_word_report(**_minimal_report_inputs(), synthesis_package=package)
    anchors, bookmarks = _anchors_and_bookmarks(report.document)

    assert anchors - bookmarks == set(), f"dangling anchors: {anchors - bookmarks}"
    assert "src_S07" not in anchors, "a label with no catalog row must not become a link"

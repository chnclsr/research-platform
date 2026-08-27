"""Deterministic DOCX report renderer for completed research runs.

The language model may supply an audited narrative, but tables, figures,
citations, and provenance are constructed from persisted run data here.  This
keeps the document export usable with the local model and prevents invented
numbers or sources from entering a report.
"""

from __future__ import annotations

import io
import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from .figure_analysis import FigureObservation, GeneratedResearchFigure
from .report_synthesis import SynthesisPackage


INK = "0B132B"
BLUE = "2563EB"
MUTED = "526175"
LINE = "CBD5E1"
HEADER_FILL = "EDF1F7"
SUCCESS = "059669"
WARNING = "D97706"
RISK = "B91C1C"
PALE_BLUE = "EAF2FF"
PALE_GREEN = "E8F5EE"
PALE_GOLD = "FFF5DD"
REPORT_PIPELINE_VERSION = "0.15.0"


@dataclass(frozen=True)
class WordReportResult:
    document: bytes
    figures: dict[str, bytes]


def _text(value: Any, limit: int | None = None) -> str:
    rendered = " ".join(str(value or "").replace("\x00", "").split())
    if limit is not None and len(rendered) > limit:
        return rendered[: max(1, limit - 1)].rstrip() + "…"
    return rendered


def _is_turkish(language: str) -> bool:
    return language.lower().startswith("tr")


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = (
        ["C:/Windows/Fonts/calibrib.ttf", "C:/Windows/Fonts/arialbd.ttf"]
        if bold
        else ["C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/arial.ttf"]
    )
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def _bar_chart(title: str, rows: Iterable[tuple[str, int]], color: str = "2563EB") -> bytes:
    data = [(str(label), int(value)) for label, value in rows if int(value) >= 0]
    width, height = 1200, max(360, 130 + 72 * max(1, len(data)))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font, label_font, value_font = _font(32, True), _font(24), _font(24, True)
    draw.text((56, 38), title, font=title_font, fill=f"#{INK}")
    if not data:
        draw.text((56, 130), "No data available", font=label_font, fill=f"#{MUTED}")
    else:
        maximum = max(value for _, value in data) or 1
        for index, (label, value) in enumerate(data):
            y = 120 + index * 72
            draw.text((56, y + 10), _text(label, 25), font=label_font, fill=f"#{INK}")
            start_x, bar_width = 390, 640
            draw.rounded_rectangle((start_x, y + 8, start_x + bar_width, y + 43), 9, fill="#EDF1F7")
            actual = max(5, int(bar_width * value / maximum)) if value else 0
            if actual:
                draw.rounded_rectangle((start_x, y + 8, start_x + actual, y + 43), 9, fill=f"#{color}")
            draw.text((1054, y + 10), str(value), font=value_font, fill=f"#{INK}")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def _metadata(source: Any) -> dict[str, Any]:
    value = getattr(source, "metadata_json", None)
    return value if isinstance(value, dict) else {}


def _publication_year(source: Any) -> str:
    metadata = _metadata(source)
    for key in ("publication_year", "year", "published_at", "publication_date", "date"):
        value = metadata.get(key)
        match = re.search(r"\b(19|20)\d{2}\b", str(value or ""))
        if match:
            return match.group(0)
    return "Bilinmiyor"


def _publication_type(source: Any, turkish: bool) -> str:
    metadata = _metadata(source)
    value = (
        metadata.get("subtype")
        or metadata.get("publication_type")
        or metadata.get("type")
        or metadata.get("document_type")
    )
    return _text(value or ("Bilinmiyor" if turkish else "Unknown"), 45)


def _source_role(source: Any, turkish: bool) -> str:
    role = str(_metadata(source).get("literature_relevance_tier") or "direct")
    if not turkish:
        return role
    return {
        "direct": "Doğrudan",
        "contextual": "Bağlamsal",
        "peripheral": "Çevresel",
    }.get(role, role.title())


def _source_evidence_counts(
    sources: list[Any],
    evidence_by_claim: dict[str, list[tuple[Any, Any]]],
) -> tuple[Counter[str], Counter[str]]:
    total: Counter[str] = Counter()
    reportable: Counter[str] = Counter()
    for links in evidence_by_claim.values():
        seen_total: set[str] = set()
        seen_reportable: set[str] = set()
        for link, source in links:
            source_id = str(source.id)
            if source_id not in seen_total:
                total[source_id] += 1
                seen_total.add(source_id)
            if (
                getattr(link, "direction", "supports") == "supports"
                and float(getattr(link, "entailment_score", 0.0) or 0.0) >= 0.5
                and source_id not in seen_reportable
            ):
                reportable[source_id] += 1
                seen_reportable.add(source_id)
    for source in sources:
        total.setdefault(str(source.id), 0)
        reportable.setdefault(str(source.id), 0)
    return total, reportable


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _style_table(
    table: Any,
    widths: list[float],
    *,
    header_fill: str = HEADER_FILL,
    font_size: float = 9,
    indent_dxa: int = 120,
) -> None:
    width_dxa = [int(round(width * 1440)) for width in widths]
    total_dxa = sum(width_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_dxa))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent_dxa))
    table_indent.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for dxa in width_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(dxa))
        grid.append(column)
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            _set_repeat_table_header(row)
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width_dxa[column_index]))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        size=font_size,
                        color=INK,
                        bold=True if row_index == 0 else None,
                    )


def _add_hyperlink(paragraph: Any, label: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def source_anchor(label: str) -> str:
    """Bookmark name for a source label. Word allows letters, digits and underscore only."""
    return f"src_{re.sub(r'[^A-Za-z0-9_]', '', label)}"[:40]


def _add_bookmarked_text(paragraph: Any, label: str, anchor: str, bookmark_id: int) -> None:
    """Write `label` into `paragraph` and wrap it in a bookmark so links can target it."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), anchor)
    paragraph._p.append(start)
    run = paragraph.add_run(label)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    return run


def _add_internal_link(paragraph: Any, label: str, anchor: str) -> None:
    """
    Same shape as _add_hyperlink but targets a bookmark in this document.

    An internal link carries `w:anchor` instead of a relationship id, so unlike the external
    variant it needs no `part.relate_to` call.
    """
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    properties.append(color)
    bold = OxmlElement("w:b")
    properties.append(bold)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    text.set(qn("xml:space"), "preserve")
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _source_link_label(source: Any, turkish: bool, limit: int = 120) -> str:
    title = _text(getattr(source, "title", ""), limit)
    if title.lower().startswith(("http://", "https://")):
        return "Kaynağı aç" if turkish else "Open source"
    return title or ("Kaynağı aç" if turkish else "Open source")


def _add_figure(
    document: Document,
    data: bytes,
    *,
    width: float,
    title: str,
    description: str,
) -> None:
    shape = document.add_picture(io.BytesIO(data), width=Inches(width))
    shape._inline.docPr.set("title", title)
    shape._inline.docPr.set("descr", description)


def _figure_width(data: bytes, *, maximum: float = 6.25, max_height: float = 6.0) -> float:
    try:
        with Image.open(io.BytesIO(data)) as image:
            aspect = image.width / max(1, image.height)
        return max(2.4, min(maximum, max_height * aspect))
    except Exception:
        return maximum


def _add_figure_interpretation(
    document: Document,
    observation: FigureObservation,
    *,
    turkish: bool,
    linkable: set[str] | None = None,
) -> None:
    main_finding = (
        " ".join(observation.main_findings[:2])
        or observation.selection_reason
    )
    limitations = " ".join(observation.limitations[:1])
    if not main_finding:
        return
    callout = document.add_table(rows=1, cols=1)
    label = "Model yorumu" if turkish else "Model interpretation"
    cell = callout.rows[0].cells[0]
    # Built run by run so the source label can carry a link to the catalog; a plain
    # cell.text assignment would collapse it into one untargetable run.
    paragraph = cell.paragraphs[0]
    paragraph.add_run(f"{label}: {main_finding} ")
    citation = f"[{observation.source_label}]"
    if observation.source_label in (linkable or set()):
        _add_internal_link(paragraph, citation, source_anchor(observation.source_label))
    else:
        paragraph.add_run(citation)
    tail = (
        f"Sınır: {limitations}"
        if turkish and limitations
        else (f"Limitation: {limitations}" if limitations else "")
    )
    if tail:
        cell.add_paragraph()
        cell.add_paragraph(tail)
    _set_cell_shading(cell, PALE_BLUE)
    _style_table(callout, [6.5], header_fill=PALE_BLUE, font_size=9.5)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11, INK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _append_field(paragraph: Any, instruction: str, display_text: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction_text, separate, display, end):
        run._r.append(node)
    _set_run_font(run, size=8, color=MUTED)


def _add_page_furniture(document: Document, run_id: str, title: str) -> None:
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_paragraph.add_run(_text(title, 90))
    _set_run_font(header_run, size=8, color=MUTED)

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"Research Platform · {run_id}  |  ")
    _set_run_font(run, size=8, color=MUTED)
    _append_field(paragraph, "PAGE")


def _add_toc_field(document: Document, turkish: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "İçindekiler alanını güncellemek için Word'de sağ tıklayın."
        if turkish
        else "Right-click and update this field in Word."
    )
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, placeholder, end):
        run._r.append(node)


def _claim_sources(claim_id: str, evidence_by_claim: dict[str, list[tuple[Any, Any]]], source_numbers: dict[str, int]) -> list[int]:
    numbers: list[int] = []
    for _, source in evidence_by_claim.get(claim_id, []):
        number = source_numbers.get(source.id)
        if number is not None and number not in numbers:
            numbers.append(number)
    return numbers


def _theme_evidence_map(
    package: SynthesisPackage,
    *,
    turkish: bool,
) -> bytes:
    sections = package.sections[:5]
    profiles = package.study_profiles[:18]
    width = 1400
    label_width = 245
    column_width = max(170, int((width - label_width - 70) / max(1, len(sections))))
    height = 190 + max(1, len(profiles)) * 48
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(32, True)
    label_font = _font(20)
    small_font = _font(17)
    draw.text(
        (42, 30),
        "Çalışmaların sentez temalarına katkısı"
        if turkish
        else "How studies contribute to synthesis themes",
        font=title_font,
        fill=f"#{INK}",
    )
    top = 118
    for column, section in enumerate(sections):
        x = label_width + column * column_width
        draw.multiline_text(
            (x + 6, top - 48),
            "\n".join(re.findall(r".{1,18}(?:\s+|$)", _text(section.title, 48))).strip(),
            font=small_font,
            fill=f"#{INK}",
            spacing=2,
        )
    for row, profile in enumerate(profiles):
        y = top + row * 48
        fill = "#F8FAFC" if row % 2 == 0 else "#FFFFFF"
        draw.rectangle((34, y, width - 34, y + 44), fill=fill)
        draw.text(
            (44, y + 11),
            f"{profile.source_label} · {_text(profile.title, 22)}",
            font=label_font,
            fill=f"#{INK}",
        )
        for column, section in enumerate(sections):
            x = label_width + column * column_width
            active = profile.source_label in section.source_ids
            cell_fill = f"#{BLUE}" if active else "#E8EEF6"
            draw.rounded_rectangle(
                (x + 12, y + 9, x + column_width - 16, y + 35),
                radius=7,
                fill=cell_fill,
            )
    if not profiles or not sections:
        draw.text(
            (44, 150),
            "Tema-kaynak eşleşmesi bulunamadı."
            if turkish
            else "No theme-to-source mapping was available.",
            font=label_font,
            fill=f"#{MUTED}",
        )
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def _add_cited_paragraph(
    document: Document, text: str, linkable: set[str] | None = None
) -> None:
    """
    Render synthesis prose, turning `[S03]` citations into links to the source catalog.

    A citation only becomes a link when the label exists in `linkable`. A link to a missing
    bookmark is silently inert in Word — the reader clicks and nothing happens, which reads
    worse than plain text.
    """
    linkable = linkable or set()
    paragraph = document.add_paragraph()
    for piece in re.split(r"(\[S\d{2,3}\])", _text(text, 12000)):
        if not piece:
            continue
        if re.fullmatch(r"\[S\d{2,3}\]", piece):
            label = piece.strip("[]")
            if label in linkable:
                _add_internal_link(paragraph, piece, source_anchor(label))
                continue
            _set_run_font(paragraph.add_run(piece), size=10.5, color=BLUE, bold=True)
        else:
            _set_run_font(paragraph.add_run(piece), size=10.5, color=INK)


def _figure_matches_section(target: str, section_title: str) -> bool:
    if not target:
        return False

    def normalize(value: str) -> set[str]:
        return {
            token
            for token in re.findall(r"[^\W\d_]{4,}", value.lower(), flags=re.UNICODE)
            if token not in {"ve", "ile", "için", "the", "and", "with"}
        }

    normalized_target = " ".join(target.lower().split())
    normalized_section = " ".join(section_title.lower().split())
    if normalized_target == normalized_section:
        return True
    if normalized_target in normalized_section or normalized_section in normalized_target:
        return True

    target_words = normalize(target)
    section_words = normalize(section_title)
    shared_words = target_words & section_words
    return (
        len(shared_words) >= 2
        and len(shared_words) / max(1, len(target_words)) >= 0.5
        and len(shared_words) / max(1, len(section_words)) >= 0.5
    )


def _add_literature_topic_map_appendix(
    document: Document,
    *,
    package: SynthesisPackage,
    figures: dict[str, bytes],
    sources: list[Any],
    turkish: bool,
    linkable_labels: set[str],
) -> None:
    """Render the complete literature topic map as Appendix B."""
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "Ek B. Literatürün Konu Haritası"
        if turkish
        else "Appendix B. Thematic Literature Landscape",
        level=1,
    )
    document.add_paragraph(
        (
            "Aşağıdaki görseller sistem performansını değil, incelenen çalışmaların neyi "
            "araştırdığını ve sentezin hangi temalarını beslediğini gösterir."
        )
        if turkish
        else (
            "The figures below describe what the included studies investigate and which synthesis "
            "themes they inform; they are not platform-performance charts."
        )
    )
    _add_figure(
        document,
        figures["16a_research_contribution_landscape.png"],
        width=6.35,
        title="Araştırma katkı türleri" if turkish else "Research contribution types",
        description=(
            "İncelenen çalışmaların odaklandığı araştırma katkılarının yatay çubuk grafiği."
            if turkish
            else "Horizontal bar chart of the research contributions addressed by included studies."
        ),
    )
    caption = document.add_paragraph(
        "Şekil B.1. Çalışmaların araştırma amacına göre literatür görünümü."
        if turkish
        else "Figure B.1. Literature landscape by study purpose."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_figure(
        document,
        figures["16b_theme_evidence_map.png"],
        width=6.35,
        title="Tema-kanıt haritası" if turkish else "Theme-evidence map",
        description=(
            "Her kaynağın rapordaki sentez temalarına yaptığı katkıyı gösteren matris."
            if turkish
            else "Matrix showing how each source contributes to the report's synthesis themes."
        ),
    )
    caption = document.add_paragraph(
        "Şekil B.2. Mavi hücre, ilgili çalışmanın o sentez temasına kanıt sağladığını gösterir."
        if turkish
        else "Figure B.2. A blue cell indicates that the study contributes evidence to that theme."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    profile_table = document.add_table(rows=1, cols=4)
    profile_headers = (
        ("Kaynak", "Araştırma katkısı", "Kanıt / tasarım", "Çalışma")
        if turkish
        else ("Source", "Research contribution", "Evidence / design", "Study")
    )
    for cell, label in zip(profile_table.rows[0].cells, profile_headers):
        cell.text = label
    source_by_id = {str(source.id): source for source in sources}
    for profile in package.study_profiles:
        row = profile_table.add_row().cells
        if profile.source_label in linkable_labels:
            _add_internal_link(
                row[0].paragraphs[0], profile.source_label, source_anchor(profile.source_label)
            )
        else:
            row[0].text = profile.source_label
        row[1].text = profile.contribution
        row[2].text = profile.evidence_design
        source = source_by_id.get(profile.source_id)
        if source is not None:
            _add_hyperlink(
                row[3].paragraphs[0],
                _source_link_label(source, turkish, 150),
                str(source.url),
            )
        else:
            row[3].text = _text(profile.title, 150)
    _style_table(profile_table, [0.55, 1.35, 1.35, 3.25], font_size=8)


def _build_synthesis_word_report(
    *,
    run_id: str,
    title: str,
    question: str,
    language: str,
    coverage: dict[str, Any],
    sources: list[Any],
    claims: list[Any],
    reportable_claims: list[Any],
    evidence_by_claim: dict[str, list[tuple[Any, Any]]],
    package: SynthesisPackage,
    scope: dict[str, Any] | None,
    sub_questions: list[str] | None,
    connector_ids: list[str] | None,
    research_mode: str,
    figure_observations: list[FigureObservation],
    research_figures: list[GeneratedResearchFigure],
) -> WordReportResult:
    """Render the synthesis-first report; retrieval diagnostics stay in appendices."""
    turkish = _is_turkish(language)
    source_numbers = {source.id: index for index, source in enumerate(sources, 1)}
    # Labels the source catalog will bookmark. Citations outside this set stay plain text
    # rather than becoming links that go nowhere.
    linkable_labels = {f"S{index:02d}" for index in source_numbers.values()}
    evidence_counts, _ = _source_evidence_counts(sources, evidence_by_claim)
    contribution_counts = Counter(profile.contribution for profile in package.study_profiles)
    figures = {
        "16a_research_contribution_landscape.png": _bar_chart(
            "Literatürde araştırılan katkı türleri"
            if turkish
            else "Research contributions represented in the literature",
            contribution_counts.most_common(),
            BLUE,
        ),
        "16b_theme_evidence_map.png": _theme_evidence_map(package, turkish=turkish),
    }

    document = Document()
    _configure_document(document)
    _add_page_furniture(document, run_id, title)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(104)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run(
        "KANITA BAĞLI SENTEZ RAPORU" if turkish else "EVIDENCE-GROUNDED SYNTHESIS REPORT"
    )
    _set_run_font(kicker_run, size=10, color=WARNING, bold=True)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(_text(title, 300))
    _set_run_font(title_run, size=28, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    subtitle_run = subtitle.add_run(
        (
            "Çalışmalar arası ortak sonuçlar, ayrışmalar ve araştırma boşlukları"
            if turkish
            else "Cross-study findings, disagreements, and research gaps"
        )
    )
    _set_run_font(subtitle_run, size=13, color=MUTED)
    cover_meta = document.add_table(rows=3, cols=2)
    cover_values = (
        (
            ("Çalışma kimliği", run_id),
            ("Üretim zamanı", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
            ("Rapor pipeline sürümü", f"v{REPORT_PIPELINE_VERSION}"),
        )
        if turkish
        else (
            ("Run ID", run_id),
            ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
            ("Report pipeline version", f"v{REPORT_PIPELINE_VERSION}"),
        )
    )
    for row, (label, value) in zip(cover_meta.rows, cover_values):
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], HEADER_FILL)
    _style_table(cover_meta, [1.6, 4.9])
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("İçindekiler" if turkish else "Contents", level=1)
    _add_toc_field(document, turkish)

    document.add_heading(
        "1. Özet" if turkish else "1. Summary",
        level=1,
    )
    lead = document.add_table(rows=1, cols=1)
    lead.rows[0].cells[0].text = _text(package.executive_summary, 6000)
    _set_cell_shading(lead.rows[0].cells[0], PALE_BLUE)
    _style_table(lead, [6.5], header_fill=PALE_BLUE, font_size=10.5)
    document.add_heading("Sonuç cümlesi" if turkish else "Bottom line", level=2)
    _add_cited_paragraph(document, package.conclusion, linkable_labels)

    document.add_heading(
        "2. Araştırma çerçevesi" if turkish else "2. Research frame",
        level=1,
    )
    document.add_heading("Ana soru" if turkish else "Primary question", level=2)
    document.add_paragraph(_text(question, 5000))
    if sub_questions:
        document.add_heading("Alt sorular" if turkish else "Sub-questions", level=2)
        for index, sub_question in enumerate(sub_questions, 1):
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{index}. ")
            _set_run_font(label, size=10.5, color=BLUE, bold=True)
            value = paragraph.add_run(_text(sub_question, 1000))
            _set_run_font(value, size=10.5, color=INK)
    scope = scope or {}
    frame = document.add_table(rows=0, cols=2)
    frame_rows = (
        (
            ("Araştırma modu", research_mode),
            (
                "Tarih aralığı",
                f"{_text(scope.get('start_date') or '—')} – {_text(scope.get('end_date') or '—')}",
            ),
        )
        if turkish
        else (
            ("Research mode", research_mode),
            (
                "Date range",
                f"{_text(scope.get('start_date') or '—')} – {_text(scope.get('end_date') or '—')}",
            ),
        )
    )
    for label, value in frame_rows:
        cells = frame.add_row().cells
        cells[0].text, cells[1].text = label, _text(value, 900)
        _set_cell_shading(cells[0], HEADER_FILL)
    _style_table(frame, [1.65, 4.85])

    document.add_heading(
        "3. Tematik kanıt sentezi" if turkish else "3. Thematic evidence synthesis",
        level=1,
    )
    if not package.sections:
        document.add_paragraph(
            "Sentez için yeterli kaynaklandırılmış bulgu bulunamadı."
            if turkish
            else "No sufficiently sourced findings were available for synthesis."
        )
    for index, section in enumerate(package.sections, 1):
        document.add_heading(f"3.{index} {_text(section.title, 240)}", level=2)
        _add_cited_paragraph(document, section.synthesis, linkable_labels)
        comparison_rows = [
            (
                "Ortak yön" if turkish else "Convergence",
                section.consensus,
                PALE_GREEN,
            ),
            (
                "Ayrışma / çelişki" if turkish else "Divergence / contradiction",
                section.disagreements,
                PALE_GOLD,
            ),
            (
                "Araştırma açısından anlamı" if turkish else "Research implication",
                section.implications,
                PALE_BLUE,
            ),
        ]
        for label, value, fill in comparison_rows:
            if not value:
                continue
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = label
            table.rows[0].cells[1].text = _text(value, 3500)
            _set_cell_shading(table.rows[0].cells[0], fill)
            _set_cell_shading(table.rows[0].cells[1], fill)
            _style_table(table, [1.55, 4.95], header_fill=fill, font_size=9.5)
        section_figures = [
            figure
            for figure in research_figures
            if _figure_matches_section(figure.section_title, section.title)
        ]
        observation_by_hash = {
            observation.image_hash: observation
            for observation in figure_observations
        }
        for figure in section_figures:
            _add_figure(
                document,
                figure.data,
                width=_figure_width(figure.data),
                title=figure.title,
                description=figure.description,
            )
            caption = document.add_paragraph(figure.caption)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(4)
            if figure.attribution:
                attribution = document.add_paragraph()
                attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = attribution.add_run(figure.attribution)
                _set_run_font(run, size=8.5, color=MUTED)
                attribution.paragraph_format.space_after = Pt(2)
            if figure.rights_statement:
                rights = document.add_paragraph()
                rights.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = rights.add_run(figure.rights_statement)
                _set_run_font(run, size=8, color=MUTED, italic=True)
                rights.paragraph_format.space_after = Pt(6)
            observation = observation_by_hash.get(figure.observation_hash)
            if observation is not None:
                _add_figure_interpretation(
                    document,
                    observation,
                    turkish=turkish,
                    linkable=linkable_labels,
                )
        section_observations = [
            observation
            for observation in figure_observations
            if _figure_matches_section(observation.recommended_section, section.title)
        ][:2]
        embedded_hashes = {figure.observation_hash for figure in section_figures}
        display_observations = [
            observation
            for observation in section_observations
            if observation.image_hash not in embedded_hashes
            and observation.main_findings
        ]
        if display_observations:
            document.add_heading(
                "Kaynak figürlerinden ek bağlam"
                if turkish
                else "Additional context from source figures",
                level=3,
            )
        for observation in display_observations:
            _add_figure_interpretation(
                document,
                observation,
                turkish=turkish,
                linkable=linkable_labels,
            )

    document.add_heading(
        "4. Çalışmalar arası değerlendirme ve sonuç"
        if turkish
        else "4. Cross-study assessment and conclusion",
        level=1,
    )
    _add_cited_paragraph(document, package.cross_study_assessment, linkable_labels)
    document.add_heading("Sonuç" if turkish else "Conclusion", level=2)
    _add_cited_paragraph(document, package.conclusion, linkable_labels)
    document.add_heading(
        "Belirsizlikler ve araştırma boşlukları"
        if turkish
        else "Uncertainties and research gaps",
        level=2,
    )
    uncertainty_box = document.add_table(rows=1, cols=1)
    uncertainty_box.rows[0].cells[0].text = _text(package.uncertainty, 6000)
    _set_cell_shading(uncertainty_box.rows[0].cells[0], PALE_GOLD)
    _style_table(uncertainty_box, [6.5], header_fill=PALE_GOLD, font_size=10)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "Ek A. Yöntem, kapsam ve yeniden üretilebilirlik"
        if turkish
        else "Appendix A. Method, coverage, and reproducibility",
        level=1,
    )
    coverage_table = document.add_table(rows=1, cols=2)
    coverage_table.rows[0].cells[0].text = "Ölçüt" if turkish else "Metric"
    coverage_table.rows[0].cells[1].text = "Değer" if turkish else "Value"
    labels = (
        ("Kaynak ailesi kapsamı", "source_family_coverage"),
        ("Sorgu dalı kapsamı", "query_branch_coverage"),
        ("İddia denetim kapsamı", "claim_audit_coverage"),
        ("Tahmini tamlık", "estimated_completeness"),
        ("Çözülmemiş ana iddia", "unresolved_major_claims"),
    ) if turkish else (
        ("Source family coverage", "source_family_coverage"),
        ("Query branch coverage", "query_branch_coverage"),
        ("Claim audit coverage", "claim_audit_coverage"),
        ("Estimated completeness", "estimated_completeness"),
        ("Unresolved major claims", "unresolved_major_claims"),
    )
    for label, key in labels:
        value = coverage.get(key)
        rendered = (
            f"{float(value):.0%}"
            if isinstance(value, float) and key != "unresolved_major_claims"
            else _text(value if value is not None else "—")
        )
        row = coverage_table.add_row().cells
        row[0].text, row[1].text = label, rendered
    _style_table(coverage_table, [4.7, 1.8])
    method_points = (
        (
            "Keşif: connector registry üzerinden çoklu akademik ve web araması.",
            "Edinim ve normalizasyon: erişim stratejisi, sürüm, hash ve kalıcı kimlik kaydı.",
            "Kanıt: pasaj konumu, alıntı, yön ve entailment puanı ile claim bağlantısı.",
            "Sentez: küçük tema paketleri; yalnız izin verilen [Sxx] kaynak kimlikleri.",
            "Sunum: konu haritası Ek B'de; retrieval ve audit ölçümleri yalnız eklerde.",
        )
        if turkish
        else (
            "Discovery: federated academic and web search through the connector registry.",
            "Acquisition and normalisation: strategy, version, hash, and persistent identifier retained.",
            "Evidence: claims link to passage location, quote, direction, and entailment score.",
            "Synthesis: bounded thematic packets with an allow-list of [Sxx] source identifiers.",
            "Presentation: topic landscape in Appendix B; retrieval and audit metrics in appendices.",
        )
    )
    for index, point in enumerate(method_points, 1):
        paragraph = document.add_paragraph()
        marker = paragraph.add_run(f"{index}. ")
        _set_run_font(marker, size=10.5, color=BLUE, bold=True)
        _set_run_font(paragraph.add_run(point), size=10.5, color=INK)
    diagnostic_text = ", ".join(
        f"{layer}={status}"
        for layer, status in package.generation_diagnostics.items()
    )
    document.add_paragraph(
        (
            f"LLM sentez kapısı: {'tam olarak geçti' if package.generated_by_llm else 'kısmen deterministik geri dönüş kullandı'}. "
            f"Katman kaydı: {diagnostic_text or 'mevcut değil'}. "
            f"Connector kapsamı: {', '.join(connector_ids or []) or 'protokol varsayılanları'}."
        )
        if turkish
        else (
            f"LLM synthesis gate: {'fully passed' if package.generated_by_llm else 'used deterministic fallback for at least one layer'}. "
            f"Layer record: {diagnostic_text or 'not available'}. "
            f"Connector scope: {', '.join(connector_ids or []) or 'protocol defaults'}."
        )
    )

    _add_literature_topic_map_appendix(
        document,
        package=package,
        figures=figures,
        sources=sources,
        turkish=turkish,
        linkable_labels=linkable_labels,
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "Ek C. Tam kaynak kataloğu" if turkish else "Appendix C. Complete source catalog",
        level=1,
    )
    source_table = document.add_table(rows=1, cols=6)
    source_headers = (
        ("#", "Yıl", "Tür", "Başlık", "Connector", "Kanıt")
        if turkish
        else ("#", "Year", "Type", "Title", "Connector", "Evidence")
    )
    for cell, label in zip(source_table.rows[0].cells, source_headers):
        cell.text = label
    for index, source in enumerate(sources, 1):
        row = source_table.add_row().cells
        label = f"S{index:02d}"
        _add_bookmarked_text(row[0].paragraphs[0], label, source_anchor(label), index)
        row[1].text = _publication_year(source)
        row[2].text = _publication_type(source, turkish)
        _add_hyperlink(
            row[3].paragraphs[0],
            _source_link_label(source, turkish, 165),
            str(source.url),
        )
        row[4].text = _text(getattr(source, "connector_id", "unknown"), 30)
        row[5].text = str(evidence_counts[str(source.id)])
    _style_table(source_table, [0.45, 0.55, 0.75, 3.45, 0.85, 0.45], font_size=7.5)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "Ek D. Denetlenmiş iddia kaydı"
        if turkish
        else "Appendix D. Audited claim register",
        level=1,
    )
    claim_table = document.add_table(rows=1, cols=5)
    claim_headers = (
        ("#", "İddia", "Durum", "Güven", "Kaynaklar")
        if turkish
        else ("#", "Claim", "Status", "Confidence", "Sources")
    )
    for cell, label in zip(claim_table.rows[0].cells, claim_headers):
        cell.text = label
    for index, claim in enumerate(reportable_claims, 1):
        row = claim_table.add_row().cells
        row[0].text = f"C{index:02d}"
        row[1].text = _text(claim.text, 280)
        row[2].text = _text(claim.status, 20)
        row[3].text = f"{float(getattr(claim, 'confidence', 0.0) or 0.0):.2f}"
        numbers = _claim_sources(str(claim.id), evidence_by_claim, source_numbers)
        row[4].text = ", ".join(f"S{number:02d}" for number in numbers) or "—"
    _style_table(claim_table, [0.45, 3.65, 0.8, 0.65, 0.95], font_size=8)

    if figure_observations:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(
            "Ek E. Kaynak figürü inceleme kaydı"
            if turkish
            else "Appendix E. Source figure observation register",
            level=1,
        )
        document.add_paragraph(
            (
                "Bu kayıt, vision modelinin kaynaklarda karşılaştığı figürlerden çıkardığı "
                "yapılandırılmış gözlemleri gösterir. Rapor anlatımına doğrudan katkı sağlayan "
                "figürler kaynak sayfasından kırpılıp ilgili bölümde gösterilir; diğerleri yalnız "
                "gözlem olarak tutulur veya güvenli biçimde yeniden çizilir."
            )
            if turkish
            else (
                "This register records structured observations extracted from source figures. "
                "Figures that materially support the report narrative are cropped from the "
                "source page and shown in context; others remain observations or are safely "
                "reconstructed."
            )
        )
        observation_table = document.add_table(rows=1, cols=7)
        headers = (
            ("Kaynak", "Sayfa", "Figür türü", "İlgi", "Raporda", "Ana bulgu", "Sınırlılık")
            if turkish
            else (
                "Source",
                "Page",
                "Figure type",
                "Relevance",
                "In report",
                "Main finding",
                "Limitation",
            )
        )
        for cell, label in zip(observation_table.rows[0].cells, headers):
            cell.text = label
        for observation in figure_observations:
            row = observation_table.add_row().cells
            if observation.source_label in linkable_labels:
                _add_internal_link(
                    row[0].paragraphs[0],
                    observation.source_label,
                    source_anchor(observation.source_label),
                )
            else:
                row[0].text = observation.source_label
            row[1].text = str(observation.page_number or "—")
            row[2].text = _text(observation.figure_type, 45)
            row[3].text = f"{observation.relevance_score:.2f}"
            row[4].text = (
                "Evet" if turkish and observation.include_in_report
                else (
                    "Hayır" if turkish
                    else ("Yes" if observation.include_in_report else "No")
                )
            )
            row[5].text = _text(
                " ".join(observation.main_findings[:2])
                or observation.selection_reason,
                320,
            )
            row[6].text = _text(
                " ".join(observation.limitations[:1]),
                220,
            )
        _style_table(
            observation_table,
            [0.5, 0.4, 0.75, 0.4, 0.55, 2.45, 1.45],
            font_size=7.5,
        )

    output = io.BytesIO()
    document.save(output)
    return WordReportResult(document=output.getvalue(), figures=figures)


def build_word_report(
    *,
    run_id: str,
    title: str,
    question: str,
    language: str,
    coverage: dict[str, Any],
    sources: list[Any],
    claims: list[Any],
    reportable_claims: list[Any],
    evidence_by_claim: dict[str, list[tuple[Any, Any]]],
    executive_summary: str,
    narrative: str,
    uncertainty: str,
    scope: dict[str, Any] | None = None,
    sub_questions: list[str] | None = None,
    connector_ids: list[str] | None = None,
    research_mode: str = "literature_scan",
    synthesis_package: SynthesisPackage | None = None,
    figure_observations: list[FigureObservation] | None = None,
    research_figures: list[GeneratedResearchFigure] | None = None,
) -> WordReportResult:
    """Build a publication-style report from audited run state.

    The model-written prose is retained, but every number, citation, source
    entry, and chart is reconstructed from persisted records.
    """
    turkish = _is_turkish(language)
    if synthesis_package is not None:
        return _build_synthesis_word_report(
            run_id=run_id,
            title=title,
            question=question,
            language=language,
            coverage=coverage,
            sources=sources,
            claims=claims,
            reportable_claims=reportable_claims,
            evidence_by_claim=evidence_by_claim,
            package=synthesis_package,
            scope=scope,
            sub_questions=sub_questions,
            connector_ids=connector_ids,
            research_mode=research_mode,
            figure_observations=figure_observations or [],
            research_figures=research_figures or [],
        )
    family_counts = Counter(str(source.family) for source in sources)
    status_counts = Counter(str(claim.status) for claim in claims)
    year_counts = Counter(_publication_year(source) for source in sources)
    connector_counts = Counter(
        str(getattr(source, "connector_id", None) or "unknown") for source in sources
    )
    evidence_counts, verified_evidence_counts = _source_evidence_counts(
        sources, evidence_by_claim
    )
    contradictory_links = sum(
        getattr(link, "direction", "") == "contradicts"
        for links in evidence_by_claim.values()
        for link, _ in links
    )
    figures = {
        "16a_source_family_distribution.png": _bar_chart(
            "Yayın yılına göre korunan literatür"
            if turkish
            else "Retained literature by publication year",
            sorted(year_counts.items(), key=lambda item: item[0]),
            BLUE,
        ),
        "16b_claim_status_distribution.png": _bar_chart(
            "İddia denetim sonucu" if turkish else "Claim audit outcome",
            sorted(status_counts.items()),
            SUCCESS,
        ),
    }

    document = Document()
    _configure_document(document)
    _add_page_furniture(document, run_id, title)
    source_numbers = {source.id: index for index, source in enumerate(sources, 1)}
    evidence_by_source: dict[str, list[tuple[Any, Any]]] = {}
    for claim in claims:
        for link, source in evidence_by_claim.get(claim.id, []):
            evidence_by_source.setdefault(str(source.id), []).append((claim, link))

    # Cover: editorial-cover pattern as a named opening override on the
    # standard_business_brief preset.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(104)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run(
        "DENETLENEBİLİR LİTERATÜR RAPORU"
        if turkish
        else "AUDITABLE LITERATURE REPORT"
    )
    _set_run_font(kicker_run, size=10, color=WARNING, bold=True)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(8)
    title_run = title_paragraph.add_run(_text(title, 300))
    _set_run_font(title_run, size=28, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    subtitle_run = subtitle.add_run(
        (
            "Kaynak keşfi, kanıt denetimi ve yeniden üretilebilirlik paketi"
            if turkish
            else "Source discovery, evidence audit, and reproducibility package"
        )
    )
    _set_run_font(subtitle_run, size=13, color=MUTED)

    cover_meta = document.add_table(rows=3, cols=2)
    cover_values = (
        (
            ("Çalışma kimliği", run_id),
            ("Üretim zamanı", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
            ("Rapor pipeline sürümü", f"v{REPORT_PIPELINE_VERSION}"),
        )
        if turkish
        else (
            ("Run ID", run_id),
            ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
            ("Report pipeline version", f"v{REPORT_PIPELINE_VERSION}"),
        )
    )
    for row, (label, value) in zip(cover_meta.rows, cover_values):
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], HEADER_FILL)
    _style_table(cover_meta, [1.6, 4.9])
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("İçindekiler" if turkish else "Contents", level=1)
    _add_toc_field(document, turkish)

    document.add_heading(
        "1. Yönetici değerlendirmesi" if turkish else "1. Executive assessment",
        level=1,
    )
    metric_table = document.add_table(rows=2, cols=4)
    metric_labels = (
        ("Korunan kaynak", "Toplam iddia", "Raporlanabilir", "Çelişen kanıt")
        if turkish
        else ("Retained sources", "All claims", "Reportable", "Contradicting evidence")
    )
    metric_values = (
        str(len(sources)),
        str(len(claims)),
        str(len(reportable_claims)),
        str(contradictory_links),
    )
    for index, label in enumerate(metric_labels):
        metric_table.rows[0].cells[index].text = label
        metric_table.rows[1].cells[index].text = metric_values[index]
        metric_table.rows[1].cells[index].paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
    _style_table(metric_table, [1.625, 1.625, 1.625, 1.625], header_fill=PALE_BLUE)

    lead = document.add_table(rows=1, cols=1)
    lead.rows[0].cells[0].text = _text(executive_summary, 5000) or (
        "Model sentezi üretilemedi; denetlenmiş bulgular aşağıdaki bölümlerde sunulmuştur."
        if turkish
        else "Model synthesis was unavailable; audited findings follow below."
    )
    _set_cell_shading(lead.rows[0].cells[0], PALE_BLUE)
    _style_table(lead, [6.5], header_fill=PALE_BLUE)

    document.add_heading(
        "2. Araştırma çerçevesi" if turkish else "2. Research frame", level=1
    )
    document.add_heading("Ana soru" if turkish else "Primary question", level=2)
    document.add_paragraph(_text(question, 5000))
    if sub_questions:
        document.add_heading(
            "Alt sorular" if turkish else "Sub-questions", level=2
        )
        for index, sub_question in enumerate(sub_questions, 1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            label = paragraph.add_run(f"{index}. ")
            _set_run_font(label, size=10.5, color=BLUE, bold=True)
            value = paragraph.add_run(_text(sub_question, 1000))
            _set_run_font(value, size=10.5, color=INK)

    frame = document.add_table(rows=0, cols=2)
    scope = scope or {}
    scope_start = _text(scope.get("start_date") or "—", 35)
    scope_end = _text(scope.get("end_date") or "—", 35)
    frame_rows = (
        (
            ("Araştırma modu", research_mode),
            ("Tarih aralığı", f"{scope_start} – {scope_end}"),
            ("Connector'lar", ", ".join(connector_ids or sorted(connector_counts))),
            ("Kaynak aileleri", ", ".join(sorted(family_counts)) or "—"),
        )
        if turkish
        else (
            ("Research mode", research_mode),
            ("Date range", f"{scope_start} – {scope_end}"),
            ("Connectors", ", ".join(connector_ids or sorted(connector_counts))),
            ("Source families", ", ".join(sorted(family_counts)) or "—"),
        )
    )
    for label, value in frame_rows:
        cells = frame.add_row().cells
        cells[0].text, cells[1].text = label, _text(value, 900)
        _set_cell_shading(cells[0], HEADER_FILL)
    _style_table(frame, [1.65, 4.85])

    document.add_heading(
        "3. Kapsam ve kanıt sağlığı" if turkish else "3. Coverage and evidence health",
        level=1,
    )
    coverage_table = document.add_table(rows=1, cols=2)
    coverage_table.rows[0].cells[0].text = "Ölçüt" if turkish else "Metric"
    coverage_table.rows[0].cells[1].text = "Değer" if turkish else "Value"
    labels = (
        ("Kaynak ailesi kapsamı", "source_family_coverage"),
        ("Sorgu dalı kapsamı", "query_branch_coverage"),
        ("İddia denetim kapsamı", "claim_audit_coverage"),
        ("Tahmini tamlık", "estimated_completeness"),
        ("Çözülmemiş ana iddia", "unresolved_major_claims"),
    ) if turkish else (
        ("Source family coverage", "source_family_coverage"),
        ("Query branch coverage", "query_branch_coverage"),
        ("Claim audit coverage", "claim_audit_coverage"),
        ("Estimated completeness", "estimated_completeness"),
        ("Unresolved major claims", "unresolved_major_claims"),
    )
    for label, key in labels:
        value = coverage.get(key)
        rendered = f"{float(value):.0%}" if isinstance(value, float) and key != "unresolved_major_claims" else _text(value if value is not None else "—")
        row = coverage_table.add_row().cells
        row[0].text, row[1].text = label, rendered
    _style_table(coverage_table, [4.7, 1.8])
    _add_figure(
        document,
        figures["16a_source_family_distribution.png"],
        width=6.35,
        title="Yayın yılı dağılımı" if turkish else "Publication year distribution",
        description=(
            "Korunan literatür kaynaklarının yayın yılına göre yatay çubuk grafiği."
            if turkish
            else "Horizontal bar chart of retained literature by publication year."
        ),
    )
    caption = document.add_paragraph(
        "Şekil 1. Yayın yılına göre korunan kaynak sayısı."
        if turkish
        else "Figure 1. Retained sources by publication year."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    _add_figure(
        document,
        figures["16b_claim_status_distribution.png"],
        width=6.35,
        title="İddia denetim sonucu" if turkish else "Claim audit outcome",
        description=(
            "Denetlenmiş iddiaların durumlarına göre yatay çubuk grafiği."
            if turkish
            else "Horizontal bar chart of audited claims by status."
        ),
    )
    caption = document.add_paragraph(
        "Şekil 2. Denetim sonrasında iddiaların dağılımı."
        if turkish
        else "Figure 2. Claim distribution after audit."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)

    document.add_heading(
        "4. Kaynaklı ana bulgular" if turkish else "4. Sourced findings", level=1
    )
    if not reportable_claims:
        document.add_paragraph(
            "Denetim kapısından geçen raporlanabilir iddia bulunamadı."
            if turkish else "No reportable claim passed the evidence gate."
        )
    else:
        for index, claim in enumerate(reportable_claims[:30], 1):
            document.add_heading(f"4.{index} {_text(claim.text, 700)}", level=2)
            status = str(claim.status)
            if turkish:
                status = {
                    "supported": "Desteklenmiş",
                    "qualified": "Koşullu / tek kaynaklı",
                    "contradicted": "Çelişkili",
                    "unresolved": "Çözülmemiş",
                }.get(status, status)
            audit = getattr(claim, "audit", {}) or {}
            finding_meta = document.add_table(rows=1, cols=3)
            finding_meta.rows[0].cells[0].text = (
                f"{'Durum' if turkish else 'Status'}\n{status}"
            )
            finding_meta.rows[0].cells[1].text = (
                f"{'Güven' if turkish else 'Confidence'}\n"
                f"{float(getattr(claim, 'confidence', 0.0) or 0.0):.2f}"
            )
            finding_meta.rows[0].cells[2].text = (
                f"{'Soru ilgisi' if turkish else 'Question relevance'}\n"
                f"{float(audit.get('question_relevance', 0.0) or 0.0):.2f}"
            )
            _style_table(
                finding_meta,
                [2.5, 1.75, 2.25],
                header_fill=PALE_GREEN if claim.status == "supported" else PALE_GOLD,
            )
            links = evidence_by_claim.get(claim.id, [])
            for link, source in links[:4]:
                quote = _text(getattr(link, "quote", ""), 450)
                location = getattr(link, "location", {}) or {}
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.first_line_indent = Inches(-0.18)
                marker = paragraph.add_run("KANIT  ")
                _set_run_font(marker, size=8.5, color=BLUE, bold=True)
                quotation = paragraph.add_run(f"“{quote}” ")
                _set_run_font(quotation, size=9.5, color=INK)
                source_number = source_numbers.get(source.id, "?")
                reference = paragraph.add_run(
                    f"[S{int(source_number):02d}] "
                    if isinstance(source_number, int)
                    else f"[S{source_number}] "
                )
                _set_run_font(reference, size=9, color=BLUE, bold=True)
                _add_hyperlink(paragraph, _source_link_label(source, turkish), source.url)
                locator = location.get("section_path") or location.get("page_number")
                if locator:
                    locator_run = paragraph.add_run(f" · {_text(locator, 80)}")
                    _set_run_font(locator_run, size=8.5, color=MUTED)
        if len(reportable_claims) > 30:
            document.add_paragraph(
                f"{len(reportable_claims) - 30} ek iddia, raporun kanıt kayıt tablosunda ve claim ledger ekinde korunmuştur."
                if turkish
                else f"{len(reportable_claims) - 30} additional claims are retained in the evidence register and claim ledger."
            )

    document.add_heading(
        "5. Kanıt sentezi" if turkish else "5. Evidence synthesis", level=1
    )
    document.add_paragraph(_text(narrative, 12000))
    document.add_heading(
        "6. Çelişkiler, boşluklar ve belirsizlik" if turkish else "6. Contradictions, gaps, and uncertainty",
        level=1,
    )
    uncertainty_box = document.add_table(rows=1, cols=1)
    uncertainty_box.rows[0].cells[0].text = _text(uncertainty, 6000) or (
        "Açık bir belirsizlik notu üretilmedi."
        if turkish
        else "No explicit uncertainty note was generated."
    )
    _set_cell_shading(uncertainty_box.rows[0].cells[0], PALE_GOLD)
    _style_table(uncertainty_box, [6.5], header_fill=PALE_GOLD)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "7. Literatür haritası" if turkish else "7. Literature landscape", level=1
    )
    connector_table = document.add_table(rows=1, cols=3)
    connector_labels = (
        ("Connector", "Korunan kaynak", "Pay")
        if turkish
        else ("Connector", "Retained sources", "Share")
    )
    for cell, label in zip(connector_table.rows[0].cells, connector_labels):
        cell.text = label
    for connector_id, count in connector_counts.most_common():
        row = connector_table.add_row().cells
        row[0].text = connector_id
        row[1].text = str(count)
        row[2].text = f"{count / max(1, len(sources)):.0%}"
    _style_table(connector_table, [3.6, 1.45, 1.45])

    document.add_heading(
        "En yüksek kanıt verimi sağlayan kaynaklar"
        if turkish
        else "Sources with the highest evidence yield",
        level=2,
    )
    yield_table = document.add_table(rows=1, cols=4)
    yield_labels = (
        ("#", "Kaynak", "Bağlı iddia", "Doğrulanmış destek")
        if turkish
        else ("#", "Source", "Linked claims", "Verified support")
    )
    for cell, label in zip(yield_table.rows[0].cells, yield_labels):
        cell.text = label
    ranked_sources = sorted(
        sources,
        key=lambda source: (
            verified_evidence_counts[str(source.id)],
            evidence_counts[str(source.id)],
        ),
        reverse=True,
    )
    for source in ranked_sources[:12]:
        row = yield_table.add_row().cells
        number = source_numbers[source.id]
        row[0].text = f"S{number:02d}"
        row[1].text = _text(source.title, 140)
        row[2].text = str(evidence_counts[str(source.id)])
        row[3].text = str(verified_evidence_counts[str(source.id)])
    _style_table(yield_table, [0.55, 4.1, 0.9, 0.95], font_size=8.5)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "8. Tam kaynak kataloğu" if turkish else "8. Complete source catalog",
        level=1,
    )
    document.add_paragraph(
        (
            "Bu tablo araştırmada kabul edilip korunan bütün kaynakları içerir. "
            "Bir kaynağın nihai sentezde az kullanılması, literatür envanterinden çıkarıldığı anlamına gelmez."
        )
        if turkish
        else (
            "This table contains every source retained by the research run. "
            "Limited use in the final synthesis does not remove a source from the literature inventory."
        )
    )
    source_table = document.add_table(rows=1, cols=7)
    source_headers = (
        ("#", "Yıl", "Tür", "Başlık", "Connector", "Rol", "Kanıt")
        if turkish
        else ("#", "Year", "Type", "Title", "Connector", "Role", "Evidence")
    )
    for cell, label in zip(source_table.rows[0].cells, source_headers):
        cell.text = label
    for index, source in enumerate(sources, 1):
        row = source_table.add_row().cells
        label = f"S{index:02d}"
        _add_bookmarked_text(row[0].paragraphs[0], label, source_anchor(label), index)
        row[1].text = _publication_year(source)
        row[2].text = _publication_type(source, turkish)
        title_paragraph = row[3].paragraphs[0]
        _add_hyperlink(title_paragraph, _source_link_label(source, turkish, 165), source.url)
        row[4].text = _text(getattr(source, "connector_id", "unknown"), 30)
        row[5].text = _source_role(source, turkish)
        row[6].text = str(evidence_counts[str(source.id)])
    _style_table(
        source_table,
        [0.42, 0.55, 0.72, 2.75, 0.85, 0.76, 0.45],
        font_size=7.5,
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(
        "9. Denetlenmiş iddia kaydı" if turkish else "9. Audited claim register",
        level=1,
    )
    claim_table = document.add_table(rows=1, cols=5)
    claim_headers = (
        ("#", "İddia", "Durum", "Güven", "Kaynaklar")
        if turkish
        else ("#", "Claim", "Status", "Confidence", "Sources")
    )
    for cell, label in zip(claim_table.rows[0].cells, claim_headers):
        cell.text = label
    for index, claim in enumerate(reportable_claims, 1):
        row = claim_table.add_row().cells
        row[0].text = f"C{index:02d}"
        row[1].text = _text(claim.text, 260)
        row[2].text = _text(claim.status, 20)
        row[3].text = f"{float(getattr(claim, 'confidence', 0.0) or 0.0):.2f}"
        numbers = _claim_sources(claim.id, evidence_by_claim, source_numbers)
        row[4].text = ", ".join(f"S{number:02d}" for number in numbers) or "—"
    _style_table(claim_table, [0.45, 3.65, 0.8, 0.65, 0.95], font_size=8)

    document.add_heading(
        "10. Yöntem ve yeniden üretilebilirlik"
        if turkish
        else "10. Method and reproducibility",
        level=1,
    )
    method_points = (
        (
            "Keşif: connector registry üzerinden çoklu akademik ve web araması.",
            "Edinim: açık içerik, AgentSearch read ve crawler fallback zinciri.",
            "Normalizasyon: URL/kalıcı kimlik/içerik hash'i ile tekilleştirme ve sürüm kaydı.",
            "Kanıt: pasaj konumu, kısa alıntı, yön ve entailment puanı ile claim bağlantısı.",
            "Sentez: yalnız denetim kapısını geçen iddialar; tam ham paket ayrıca korunur.",
        )
        if turkish
        else (
            "Discovery: federated academic and web search through the connector registry.",
            "Acquisition: open content, AgentSearch read, and crawler fallback chain.",
            "Normalization: persistent identifier, canonical URL, and content-hash deduplication.",
            "Evidence: claim links retain passage locator, quote, direction, and entailment score.",
            "Synthesis: only audited claims enter prose; the complete raw package is retained separately.",
        )
    )
    for index, point in enumerate(method_points, 1):
        paragraph = document.add_paragraph()
        number_run = paragraph.add_run(f"{index}. ")
        _set_run_font(number_run, size=10.5, color=BLUE, bold=True)
        text_run = paragraph.add_run(point)
        _set_run_font(text_run, size=10.5, color=INK)

    document.add_paragraph(
        (
            "Bu Word raporu okunabilir ana teslimattır; claim ledger, evidence matrix, "
            "ham kaynaklar, ham pasajlar ve yeniden üretilebilirlik manifesti ZIP paketinde ayrıca korunur."
        )
        if turkish
        else (
            "This Word report is the readable primary deliverable; the claim ledger, evidence "
            "matrix, raw sources, raw passages, and reproducibility manifest remain in the ZIP bundle."
        )
    )

    output = io.BytesIO()
    document.save(output)
    return WordReportResult(document=output.getvalue(), figures=figures)

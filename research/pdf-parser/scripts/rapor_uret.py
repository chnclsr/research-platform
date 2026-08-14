"""RAPOR.docx uretir.  Kullanim: .\.venv\Scripts\python.exe rapor_uret.py"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

d = Document()

# ---------- stil ----------
st = d.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(8)
st.paragraph_format.line_spacing = 1.15

for ad, boy in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    s = d.styles[ad]
    s.font.name = "Calibri"
    s.font.size = Pt(boy)
    s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    s.paragraph_format.space_before = Pt(16)
    s.paragraph_format.space_after = Pt(6)

for sec in d.sections:
    sec.left_margin = sec.right_margin = Cm(2.4)
    sec.top_margin = sec.bottom_margin = Cm(2.2)


def P(metin="", stil=None, kalin=False, italik=False, boyut=None):
    p = d.add_paragraph(style=stil)
    r = p.add_run(metin)
    r.bold, r.italic = kalin, italik
    if boyut:
        r.font.size = Pt(boyut)
    return p


def zengin(parcalar):
    """[(metin, kalin), ...] seklinde karisik bicimli paragraf."""
    p = d.add_paragraph()
    for metin, kalin in parcalar:
        r = p.add_run(metin)
        r.bold = kalin
    return p


def madde(metin, seviye=0):
    p = d.add_paragraph(metin, style="List Bullet" if seviye == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    return p


def tablo(basliklar, satirlar, genislikler=None):
    t = d.add_table(rows=1, cols=len(basliklar))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, b in enumerate(basliklar):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(b)
        r.bold = True
        r.font.size = Pt(9.5)
    for sat in satirlar:
        cells = t.add_row().cells
        for i, v in enumerate(sat):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    for row in t.rows:
        for c in row.cells:
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if genislikler:
        for row in t.rows:
            for i, w in enumerate(genislikler):
                row.cells[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


SEKIL = [0]


def gorsel(dosya, aciklama, genislik_cm=16.0):
    """Ortalanmis gorsel + altina numarali aciklama."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(f"gorseller/{dosya}", width=Cm(genislik_cm))

    SEKIL[0] += 1
    a = d.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_after = Pt(12)
    r = a.add_run(f"Şekil {SEKIL[0]} — {aciklama}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return a


def kod(metin):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(metin)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ================= KAPAK =================
b = d.add_paragraph()
b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = b.add_run("PDF Metin Çıkarma Kütüphaneleri Karşılaştırması")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

b = d.add_paragraph()
b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = b.add_run("pypdf ile pdf-inspector'ın araştırma platformu bağlamında değerlendirilmesi")
r.font.size = Pt(13)
r.italic = True

b = d.add_paragraph()
b.alignment = WD_ALIGN_PARAGRAPH.CENTER
b.paragraph_format.space_before = Pt(18)
r = b.add_run("13 Ağustos 2026")
r.font.size = Pt(10.5)

d.add_paragraph()

# ================= 2. NEDEN =================
d.add_heading("1. Bu değerlendirme neden yapıldı", level=1)

P("Platform bir PDF'i indirdikten sonra metnini çıkarıyor, pasajlara bölüyor ve bu pasajları "
  "kanıt olarak kullanıyor. Metin çıkarma aşamasındaki bir hata sonraki her aşamaya taşınıyor. "
  "Şu anki kod, her sayfanın metnini alıp üstüne yapay bir sayfa başlığı ekleyerek birleştiriyor:")

kod('reader = PdfReader(io.BytesIO(response.content))\n'
    'text = "\\n\\n".join(\n'
    '    f"# Page {index}\\n\\n{page.extract_text() or \'\'}"\n'
    '    for index, page in enumerate(reader.pages, start=1)\n'
    ')')

P("Buradaki kritik nokta şu: sistem, metindeki başlıkları kullanarak her pasaja bir bölüm yolu "
  "atıyor ve bu yola bakarak pasajın kanıt sayılıp sayılmayacağına karar veriyor. Örneğin "
  "kaynakça bölümünden gelen bir cümle kanıt olarak kabul edilmemeli. Ne var ki pypdf hiçbir "
  "gerçek bölüm başlığı üretmiyor — çıktıdaki tek başlık, kodun kendi eklediği “# Page 1”, "
  "“# Page 2” satırları. Dolayısıyla bu kontrol mekanizması PDF kaynakları için hiç devreye "
  "girmiyor. Test ettiğimiz dokuz dosyanın dokuzunda da böyle.")

P("İkinci gerekçe, kullanıcılardan gelen “paragraflar karışıyor” şikâyeti. Bunun gerçekten iki "
  "sütunlu sayfalarda okuma sırasının bozulmasından mı, yoksa başka bir şeyden mi kaynaklandığı "
  "belirsizdi. Bu da incelendi.")

# ================= 3. YÖNTEM =================
d.add_heading("2. Nasıl test edildi", level=1)

d.add_heading("2.1 Ölçme yaklaşımı", level=2)

P("Her dosya iki kütüphaneden de geçirildi ve çıktılar karşılaştırıldı. Hız ölçümlerinde her "
  "dosya üç kez çalıştırıldı, ısınma koşusu atıldı, medyan alındı. Sürümler: pdf-inspector "
  "1.14.1, pypdf 6.15.0, Python 3.10, Windows 11.")

P("Otomatik sayımlar bir yere kadar götürüyor: bir sayaç “şu desen kaç kez geçti” diyebilir, "
  "ama o geçişin gerçekten hata olup olmadığını söyleyemez. Bu yüzden dört belge sayfa sayfa "
  "elle okundu. Bu sırada birkaç otomatik ölçümün yanıltıcı olduğu görüldü ve rapordan "
  "çıkarıldı.")

d.add_heading("2.2 “Doğru” neye göre belirlendi", level=2)

P("İki çıktıyı karşılaştırırken hangisinin doğru olduğuna karar vermek için bir ölçüt gerekiyor. "
  "Burada ölçüt, çıktıyı tüketen sistemin ihtiyaçlarından türetildi. Sistem dört şey istiyor:")

madde("Cümleler bütün olsun — bir pasaj kendi içinde anlaşılabilir olmalı.")
madde("Alıntılar birebir bulunabilsin — bir cümle metinde aynen aranıp bulunabilmeli.")
madde("Bölüm bilgisi olsun — pasajın hangi bölümden geldiği bilinmeli.")
madde("Sayılar etiketlerinden kopmasın — bir değerin hangi satır ve sütuna ait olduğu belli olmalı.")

P("Bir çıktı bu dördünden birini bozuyorsa hata sayıldı. Sadece daha çirkin görünüyorsa "
  "sayılmadı. Örneğin fazladan boş satır, girinti farkı veya kalın/italik biçimlendirme "
  "hata olarak değerlendirilmedi.")

d.add_heading("2.3 Test edilen belgeler ve seçim gerekçeleri", level=2)

P("Belgeler rastgele seçilmedi. Her biri belirli bir zorluğu sınamak için alındı; böylece bir "
  "kütüphanenin nerede iyi, nerede kötü çalıştığı görülebilsin diye.")

tablo(
    ["Belge", "Sayfa", "Neden seçildi"],
    [
        ["BERT makalesi", "16", "İki sütunlu düzen ve sayfa altı dipnotları — en yaygın akademik biçim"],
        ["Türkçe makale", "6", "Türkçe karakter kümesi, iki sütun, yan yana duran tablolar"],
        ["Sybil (tıp)", "17", "Tıp dergisi düzeni: kutulu bilgi alanları, her sayfada künye, yoğun kaynakça"],
        ["VGG makalesi", "14", "Tablo ağırlıklı, tek sütun — tablo kurma yeteneğini sınamak için"],
        ["ResNet makalesi", "12", "İki sütun ve görsel ağırlıklı"],
        ["Attention makalesi", "15", "Karmaşık tablolar ve şekiller"],
        ["GPT-3 raporu", "75", "Uzun belge — dayanıklılık"],
        ["GPT-4 raporu", "100", "Uzun ve görsel ağırlıklı; ayrıca iki belgenin birleştirilmiş hâli"],
        ["Taranmış BERT", "6", "Metin katmanı olmayan taranmış belge"],
    ],
    [3.6, 1.4, 11.0])

P("Dördü sayfa sayfa elle okundu: BERT, Türkçe makale, Sybil ve VGG. Diğer beşi otomatik "
  "ölçümlere dahil edildi. Bu dört belgenin seçilmesinin nedeni, elle inceleme sırasında ortaya "
  "çıkan her kusur türünün en az iki farklı belgede tekrarlanmış olması; yani gözlemler tek "
  "belgeye özgü tesadüfler değil.", italik=True)

# ================= 4. HATA TÜRLERİ =================
d.add_heading("3. Aranan hata türleri", level=1)

P("Rapor boyunca kullanılan terimler aşağıda açıklanıyor. Her hata türü için ne olduğu, nasıl "
  "göründüğü ve sistemi neden ilgilendirdiği ayrı ayrı yazıldı.")

d.add_heading("3.1 Cümle bütünlüğünü bozan hatalar", level=2)

zengin([("Okuma sırası bozukluğu. ", True),
        ("İki sütunlu bir sayfada metin, soldaki sütunun tamamı bitmeden sağdakine atlarsa "
         "cümleler birbirine karışır. Sonuçta ortaya çıkan metin dilbilgisi açısından anlamsızdır "
         "ama gözle bakınca “bir şeyler yazıyor” izlenimi verir. Sistemde bu, birbirine "
         "yapışmış iki ayrı konudan oluşan pasajlar üretir.", False)])

zengin([("Sayfa künyesinin metne karışması. ", True),
        ("Sayfanın üstünde ve altında tekrar eden, içerikle ilgisi olmayan satırlar: dergi adı, "
         "sayfa numarası, telif satırı, “şu tarihte şuradan indirildi” damgası. Alan yazında "
         "bunlara topluca ", False),
        ("page furniture", True),
        (" denir. Kendi satırlarında durdukları sürece zararsızdırlar; sorun bir cümlenin "
         "ortasına düştüklerinde başlar. O zaman cümle ikiye bölünür ve alıntı doğrulama o "
         "cümleyi metinde bulamaz.", False)])

zengin([("Dipnotun gövdeye karışması. ", True),
        ("Sayfa altındaki dipnot metninin normal paragraf akışına girmesi. Dipnot kendi başına "
         "değerli bilgi olabilir, ama gövde cümlesinin ortasına girdiğinde iki metin birbirine "
         "karışır.", False)])

d.add_heading("3.2 Yapı bilgisini etkileyen hatalar", level=2)

zengin([("Kaçırılan başlık. ", True),
        ("Belgede gerçekten başlık olan bir satırın düz metin olarak çıkması. Sistem bölüm "
         "yolunu başlıklardan türettiği için, kaçırılan bir “Kaynakça” başlığı, kaynakça "
         "cümlelerinin kanıt olarak kabul edilmesine yol açar.", False)])

zengin([("Uydurma başlık. ", True),
        ("Başlık olmayan bir metnin başlık olarak işaretlenmesi. Yazar adları, şekil "
         "etiketleri, kenar damgaları veya yarım kalmış cümleler bu şekilde başlığa "
         "dönüşebiliyor. Sonuç, yanlış bölüm yolları.", False)])

zengin([("Başlığın gövdeye yapışması. ", True),
        ("Başlığın kendi satırında kalmayıp altındaki paragrafın ilk cümlesine eklenmesi. "
         "Bu, kaçırılan başlıktan daha kötüdür: başlık kendi satırında dursaydı sonradan basit "
         "bir kuralla yakalanabilirdi, yapıştığında o imkân da kalmaz.", False)])

d.add_heading("3.3 Tablolarla ilgili hatalar", level=2)

zengin([("Tablo yapısının bozulması. ", True),
        ("Bir sayının hangi satır ve sütuna ait olduğunun kaybolması. İki farklı biçimde ortaya "
         "çıkıyor: ya tablo hiç kurulmuyor ve sayılar düz metne dökülüyor, ya da tablo kuruluyor "
         "ama hücreler yanlış eşleşiyor. İkincisi daha risklidir, çünkü çıktı düzgün bir tablo "
         "gibi görünür.", False)])

zengin([("Sütun başlığının düşmesi. ", True),
        ("Tablo kuruluyor, sayılar doğru sırada duruyor, ama sütun adları çıktıya girmemiş. "
         "Değerler orada ama neyi ölçtükleri yazmıyor.", False)])

d.add_heading("3.4 Metin düzeyindeki hatalar", level=2)

zengin([("Bağlı harf (ligatür) kalıntısı. ", True),
        ("Basılı metinde “fi” ve “fl” gibi harf çiftleri tek bir karakter olarak dizilir. Bu "
         "karakter normal harflere çevrilmezse “official” kelimesi arama sırasında bulunamaz, "
         "çünkü içindeki “fi” aslında farklı bir karakterdir.", False)])

zengin([("Birleşmemiş tireleme. ", True),
        ("Satır sonunda bölünen kelimenin tekrar birleştirilmemesi (“evalu- ation”). "
         "Alıntı doğrulamayı bozar.", False)])

zengin([("Metin içi göndermelerde boşluğun kaybolması. ", True),
        ("Makale içinde başka bir şekle veya bölüme yapılan atıfların birbirine yapışması: "
         "“Figure 1” ifadesinin “Figure1” hâline gelmesi gibi. Okunabilirlik açısından küçük, "
         "arama açısından rahatsız edici bir kusur.", False)])

P("Hata sayılmayanlar: fazladan boş satır, girinti farkları, madde işareti stili, Markdown "
  "biçimlendirme tercihleri ve taranmış bir belgeden metin çıkmaması.", italik=True)

# ================= 5. ÖLÇÜM SONUÇLARI =================
d.add_heading("4. Ölçüm sonuçları", level=1)

d.add_heading("4.1 Hız", level=2)

P("pdf-inspector Rust ile yazılmış ve bu farkı ölçümlerde açıkça gösteriyor. Dokuz dosyanın "
  "hepsinde daha hızlı; oran altı ila on bir kat arasında değişiyor.")

tablo(
    ["Belge", "Boyut", "pypdf", "pdf-inspector", "Oran"],
    [
        ["VGG makalesi", "0,20 MB", "0,477 sn", "0,042 sn", "11,4×"],
        ["Attention makalesi", "2,22 MB", "0,538 sn", "0,059 sn", "9,1×"],
        ["GPT-4 raporu", "5,25 MB", "1,187 sn", "0,132 sn", "9,0×"],
        ["ResNet makalesi", "0,82 MB", "0,302 sn", "0,039 sn", "7,7×"],
        ["GPT-3 raporu", "6,77 MB", "0,847 sn", "0,111 sn", "7,6×"],
        ["BERT makalesi", "0,78 MB", "0,381 sn", "0,054 sn", "7,1×"],
        ["Sybil makalesi", "0,80 MB", "0,272 sn", "0,039 sn", "7,0×"],
        ["Türkçe makale", "0,12 MB", "0,133 sn", "0,021 sn", "6,3×"],
    ],
    [4.6, 2.0, 2.2, 3.0, 1.8])

P("Yine de pypdf'in en yavaş sonucu bile yüz sayfalık bir belge için 1,2 saniye; yani hız "
  "mevcut iş hacminde darboğaz değil. Fark gerçek, ama tek başına geçiş gerekçesi sayılmamalı. "
  "Hiçbir dosyada çökme veya bellek sorunu görülmedi. Taranmış belgeyi pdf-inspector doğru "
  "sınıflandırdı (0,95 güven) ve altı sayfanın altısı için OCR gerektiğini bildirdi; pypdf'in "
  "böyle bir sinyali yok.")

d.add_heading("4.2 İçerik kaybı", level=2)

P("İki kütüphane de belgelerin tamamını okudu; hiçbir dosyada sayfa atlanmadı veya bölüm "
  "kayboldu. Karakter sayısı farkı en fazla yüzde üç ve bu fark boşluk ile biçimlendirme "
  "tercihlerinden geliyor. Yani pdf-inspector'ın kusurları içeriği kaybetmekle değil, içeriği "
  "yanlış düzenlemekle ilgili.")

d.add_heading("4.3 Bölüm başlıkları ve kaynakça kontrolü", level=2)

P("Bu, iki kütüphane arasındaki en büyük işlevsel fark. pypdf hiçbir belgede gerçek bölüm "
  "başlığı üretmiyor. pdf-inspector dokuz belgenin yedisinde kaynakça başlığını tanıdı ve "
  "sistemdeki kaynakça filtresini çalışır hâle getirdi.")

P("Aşağıdaki tablo, belgelerdeki numaralı bölüm başlıklarının (“2.1 Architecture” gibi) "
  "pdf-inspector çıktısında gerçekten başlık olarak işaretlenip işaretlenmediğini gösteriyor. "
  "Sayımdan önce içindekiler satırları ve numaralı liste öğeleri elendi.")

tablo(
    ["Belge", "Bölüm başlığı", "Tanınan", "Oran"],
    [
        ["Attention makalesi", "22", "22", "%100"],
        ["GPT-3 raporu", "39", "39", "%100"],
        ["BERT makalesi", "18", "18", "%100"],
        ["ResNet makalesi", "11", "11", "%100"],
        ["GPT-4 raporu", "10", "10", "%100"],
        ["VGG makalesi", "16", "5", "%31"],
    ],
    [5.0, 3.6, 2.8, 2.4])

P("Numaralı başlıklarda pdf-inspector altı belgenin beşinde kusursuz. Tek istisna VGG "
  "makalesi: orada bütün üst düzey başlıkları (1, 2, 3, 4, 5) doğru bulmuş, ama on bir alt "
  "bölüm başlığının hepsini kaçırıp altındaki paragrafa yapıştırmış.")

P("Numarasız başlıklarda ise sonuç bambaşka. Sybil makalesi bölümlerini büyük harfle "
  "adlandırıyor (“MATERIALS AND METHODS”, “RESULTS”, “REFERENCES”); on iki başlıktan yalnızca "
  "biri tanınmış. Kaynakça başlığı tanınmadığı için bu belgede filtre hiç devreye girmemiş. "
  "Dikkat çekici olan şu: tanınan başlık ile kaçırılanlar PDF'te aynı puntoda ve aynı biçimde "
  "dizilmiş. Yani hangi başlığın tanınacağı belgenin görünüşünden kestirilemiyor.")

P("Bir nokta daha var. pypdf o on iki başlığın hepsini kendi satırında bırakıyor; "
  "pdf-inspector dördünü paragrafa yapıştırıyor, ikisini paragraf ortasına gömüyor. Yani basit "
  "bir sonradan-düzeltme kuralı pypdf çıktısından on iki başlığı da kurtarabilirken, "
  "pdf-inspector çıktısından yalnızca beşini kurtarabilir.")

d.add_heading("4.4 Tablolar", level=2)

P("pypdf tablo kavramını hiç tanımıyor, metni düz satırlar hâlinde döküyor. pdf-inspector ise "
  "Markdown tabloları üretiyor. İlk bakışta net bir üstünlük gibi görünüyor ve bazı belgelerde "
  "gerçekten öyle: BERT makalesindeki sonuç tablosunda pdf-inspector sayı-etiket ilişkisini "
  "korurken pypdf tamamen dağıtıyor. Ne var ki üretilen 92 tablo bloğunun yapısal tutarlılığı "
  "ölçüldüğünde durum değişiyor:")

tablo(
    ["Belge", "Tablo bloğu", "Sütun hizası bozuk", "Sütun adı kayıp"],
    [
        ["Sybil makalesi", "24", "12", "2"],
        ["GPT-3 raporu", "17", "11", "1"],
        ["BERT makalesi", "16", "11", "6"],
        ["GPT-4 raporu", "14", "13", "2"],
        ["VGG makalesi", "9", "7", "2"],
        ["ResNet makalesi", "6", "3", "3"],
        ["Attention makalesi", "4", "3", "2"],
        ["Türkçe makale", "2", "2", "2"],
        ["TOPLAM", "92", "62 (%67)", "20 (%22)"],
    ],
    [4.6, 3.0, 4.2, 3.4])

P("Blokların üçte ikisinde başlık satırının ilan ettiği sütun sayısı ile veri satırlarının "
  "hücre sayısı uyuşmuyor; beşte birinde en az bir sütun adı düşmüş. Kusur dokuz belgenin "
  "dokuzunda da var, yani belgeye özgü değil.")

P("Somut bir örnek Sybil makalesinin ana sonuç tablosu. Orijinali aşağıda: üç veri seti satırı "
  "ve yedi sütun, her değer kendi satır ve sütununda.")

gorsel("01_sybil_tablo1.png",
       "Sybil makalesi, sayfa 5. Orijinal TABLE 1 — üç veri seti, yedi sütun.", 16.2)

P("pdf-inspector bu tabloyu üç ayrı parçaya bölmüş. Başlık satırı dört sütun ilan ederken veri "
  "satırlarında altı hücre var; birinci ve ikinci yıl değerleri aynı hücreye sıkışmış. Üçüncü "
  "parçada ise satır etiketleri tamamen kaybolmuş:")

kod("|Data Set 1-Year Risk...|2-Year Risk...|3-Year Risk...|4-Year Risk... 5-Year Risk...|\n"
    "|NLST|0.92 (0.88 to 0.95) 0.86 (0.82 to 0.90)|0.80 (0.77 to 0.84)|0.77 ... 0.75 ...|\n"
    "                       ^ birinci ve ikinci yil ayni hucrede\n\n"
    "||6-Year Risk, AUC (95% CI)|C-Index (95% CI)|\n"
    "||0.75 (0.72 to 0.78) NA 0.74 (0.66 to 0.81)|0.75 (0.72 to 0.78) 0.81 (...) 0.80 (...)|\n"
    "   ^ uc veri setinin degeri tek hucrede, NLST/MGH/CGMH etiketleri yok")

P("Sonuçta altı yıllık başarı değerlerinin hangi veri setine ait olduğu çıkarılamıyor. "
  "Yukarıdaki blok yine de geçerli bir Markdown tablosu; hatalı olduğu ancak orijinalle "
  "karşılaştırılınca anlaşılıyor.")

P("VGG makalesinde durum daha hafif ama aynı yönde: değerlerin sırası doğru, sütun başlıkları "
  "(“top-1 hata oranı”, “top-5 hata oranı”) düşmüş — sayılar orada, neyi ölçtükleri yazmıyor. "
  "Aynı belgede beklenmedik bir davranış daha var: normal akan bir paragraf keyfi kelime "
  "sınırlarından hücrelere bölünerek tabloya çevrilmiş, oysa orijinal sayfada o bölgede tabloya "
  "benzeyen hiçbir öge yok. Korpusta bu türden beş vaka bulundu.")

kod("pdf-inspector:\n"
    "  |are significantly|less|deep than|ours, and|they did|not evaluate|on the large-scale|\n"
    "  |dataset.|Goodfellow et al.|(2014)|applied|deep ConvNets|(11|weight layers)|\n\n"
    "pypdf:\n"
    "  are significantly less deep than ours, and they did not evaluate on the large-scale\n"
    "  ILSVRC dataset. Goodfellow et al. (2014) applied deep ConvNets (11 weight layers)")

zengin([("Bu bulgunun özü şu: ", True),
        ("pypdf tablo konusunda yapısız ama dürüst bir çıktı veriyor — sayıların okuma sırası "
         "korunuyor ve çıktının eksik olduğu bellidir. pdf-inspector kendinden emin ama "
         "kısmen yanlış bir yapı veriyor. Geçerli bir Markdown tablosu gibi göründüğü için "
         "hatanın fark edilme ihtimali düşük. Kanıta dayalı bir sistemde bir başarı değerinin "
         "yanlış veri setine atfedilmesi, o değerin hiç yapılandırılmamış olmasından daha "
         "risklidir.", False)])

d.add_heading("4.5 Metin düzeyindeki farklar", level=2)

P("pdf-inspector'ın en tartışmasız üstünlüğü bağlı harflerde. Test edilen belgelerde pypdf "
  "63 ile 444 arasında değişen sayıda çevrilmemiş bağlı harf bıraktı; pdf-inspector "
  "istisnasız sıfır. Bu doğrudan alıntı doğrulamayı etkileyen bir kazanç.")

P("Buna karşılık pdf-inspector çapraz referanslarda boşluk düşürüyor. BERT makalesinde "
  "“Section” kelimesinin sekiz geçişinin sekizinde de sonraki numaraya yapışmış "
  "(“Section5.1that”); pypdf'de bu sorun yok. İki davranış aynı mekanizmadan doğuyor: "
  "pdf-inspector satırları birleştirip akıcı metin ürettiği için hem daha temiz bir sonuç "
  "veriyor hem de birleştirme noktasında boşluk kaybediyor.")

P("Satır sonu tirelemesini ikisi de çözmüyor, sadece farklı biçimde bırakıyor. pypdf satır "
  "sonunu koruyup tireyi orada bırakıyor, pdf-inspector satırları birleştirip tireyi metnin "
  "içinde bırakıyor. Sonuç ikisinde de bozuk.")

P("Dipnot işaretlerinde pdf-inspector daha doğru davranıyor: gerçek üst simge karakterleri "
  "üretiyor ve işareti ait olduğu yerde bırakıyor. pypdf işareti çoğu zaman sonraki cümlenin "
  "başına kaydırıyor.")

d.add_heading("4.6 Hangi fonksiyonun kullanıldığı sonucu değiştiriyor", level=2)

P("pdf-inspector üç ayrı metin çıkarma yolu sunuyor ve bunlar aynı belgede farklı sonuç veriyor. "
  "Bu raporun tüm ölçümleri process_pdf ile alındı. Adı nedeniyle pypdf'in karşılığı gibi duran "
  "extract_text fonksiyonu ise üç belgede ciddi sorun çıkardı: Sybil ve VGG'de kelime "
  "aralarındaki boşlukları tamamen kaybetti (“originalreports Sybil:AValidatedDeepLearningModelto”), "
  "Türkçe makalede Türkçe karakterleri sildi ve yüz sayfalık GPT-4 raporunda içeriğin yaklaşık "
  "yüzde altmışını hiçbir uyarı vermeden atladı. Ayrıca sayfa bazlı API sayfaları sıfırdan "
  "numaralandırıyor, pypdf ve PDF okuyucular birden.")

# ================= 6. BELGE BELGE =================
d.add_heading("5. Belge belge gözlemler", level=1)

P("Bu bölüm, elle okunan dört belgede neyin nasıl çıktığını anlatıyor. Sıralama, "
  "pdf-inspector'ın en iyi sonuç verdiği belgeden en kötü sonuç verdiğine doğru.")

d.add_heading("5.1 BERT makalesi — pdf-inspector belirgin biçimde önde", level=2)

P("İki sütunlu, dipnotlu, orta uzunlukta bir konferans makalesi. Beklenen zorluk sütun okuma "
  "sırasıydı; öyle çıkmadı. Her iki kütüphane de cümleleri bütün bıraktı, sütun karışması "
  "görülmedi. Yani “paragraflar karışıyor” şikâyetinin bu belgede karşılığı yok.")

P("pdf-inspector burada üç yerde açık ara kazandı: 157 bağlı harfi temizledi, sonuç tablosunun "
  "yapısını korudu ve kırk üç başlığın otuz beşini doğru buldu — kaynakça filtresini çalıştıran "
  "başlık dahil. pypdf bu belgede tek bir gerçek bölüm başlığı üretmedi. Kusur tarafında ise "
  "“Figure” ve “Table” kelimelerinin yarısında, “Section” kelimesinin tamamında boşluk düşmüş; "
  "ayrıca iki satıra taşan üç başlığı ikiye bölmüş ve arXiv kenar damgasını başlık saymış.")

P("İki kütüphanenin ortak başarısızlığı dipnotlarda. Hiçbiri dipnot metnini gövdeden ayırmıyor; "
  "ikisi de sayfa bloğunun sonuna koyuyor. Sayfalar birleştirilince dipnot metni iki paragrafın "
  "arasına giriyor. Bu, kullanıcı şikâyetinin muhtemel gerçek kaynağı — sütun karışması değil, "
  "dipnotun gövdeye karışması.")

d.add_heading("5.2 VGG makalesi — karışık sonuç", level=2)

P("Tek sütunlu, tablo ağırlıklı bir makale. Her sayfanın üstünde tekrar eden bir konferans "
  "damgası var. pypdf bu damgayı on dört kez, her sayfada bir kez metne alıyor; pdf-inspector "
  "yalnızca bir kez. Platform açısından pdf-inspector'ın davranışı tercih edilir, çünkü aksi "
  "hâlde on dört gürültü satırı pasajlara karışır. Ancak bu temizlik tutarlı değil — Sybil "
  "makalesinde aynı türden bir damga on yedi sayfanın on yedisinde de duruyor.")

P("Buna karşılık başlıklarda belirgin bir kayıp var. Makalede on üç numaralı alt bölüm başlığı "
  "bulunuyor ve pdf-inspector hiçbirini başlık olarak çıkarmamış; hepsini kendi paragrafının "
  "ilk cümlesine yapıştırmış. pypdf bunları kendi satırlarında tutuyor.")

gorsel("04_vgg_baslik.png",
       "VGG makalesi, sayfa 2. “2.1 ARCHITECTURE” orijinalde kendi satırında duran bir başlık.",
       15.5)

kod("pdf-inspector:  2.1 ARCHITECTURE During training, the input to our ConvNets is a...\n\n"
    "pypdf:          2.1 ARCHITECTURE\n"
    "                During training, the input to our ConvNets is a fixed-size...")

P("Tablolarda ise sonuç orta hâlli: sütunlar karışmamış, değerlerin sırası doğru, ama sütun "
  "başlıkları düşmüş. Ayrıca daha önce bahsedilen “paragrafın tabloya çevrilmesi” vakası bu "
  "belgede görüldü.")

d.add_heading("5.3 Türkçe makale — başa baş", level=2)

P("İki sütunlu bir dergi makalesi. Türkçe karakterler her iki kütüphanede de sorunsuz çıktı; "
  "1765 Türkçe karakterin tamamı korundu, bozuk kodlama uyarısı üretilmedi.")

P("Asıl bulgu okuma sırasında. pdf-inspector altı sayfanın üçünde sütunları satır satır "
  "birbirine geçirdi; pypdf hepsinde doğru okudu. Sayfada sütunlar görsel olarak apaçık ayrık, "
  "yani belirsiz bir durum değil:")

gorsel("02_turkce_s4_sutunlar.png",
       "Türkçe makale, sayfa 4. İki sütun ve yan yana duran üç tablo.", 10.5)

P("Sayfanın alt kısmındaki gövde metni, hatanın oluştuğu yer. Orijinalde iki sütun ayrı ayrı "
  "akıyor:")

gorsel("02b_turkce_s4_govde.png",
       "Aynı sayfanın alt kısmı — sol ve sağ sütun birbirinden bağımsız iki paragraf.", 16.2)

P("pdf-inspector bu iki sütunu satır satır iç içe geçirmiş. Aşağıdaki çıktıda soldaki ve "
  "sağdaki paragrafın cümleleri dönüşümlü olarak sıralanıyor:")

kod("pdf-inspector (sayfa 4):\n"
    "  Tablo 6'da Katılımcıların hakemlik kademelerine göre .05 anlamlı- Tarafından yapılan\n"
    "  çalışmada ise futbol hakemlerinde cinsiyete lık düzeyinde antrenörlüğe ilişkin karar\n"
    "  verme stillerinin belirlen- göre karar verme esnasında bocalama, sorumluluğu başkasına")

P("Aynı sorun kaynakça sayfasında da tekrarladı; bir kaynağın adı ikiye bölünüp arasına başka "
  "bir kaynak girdi.")

P("Tablolarda da BERT'teki durumun tersi yaşandı. Sayfada yan yana duran iki tabloyu "
  "pdf-inspector tek bir ızgaraya birleştirdi; tablo başlıkları bile iç içe girdi ve hangi "
  "sayının hangi tabloya ait olduğu kayboldu. Bu belgede pypdf'in düz çıktısı daha okunaklı.")

P("pdf-inspector'ın bu belgedeki net kazancı sayfa numaralarında. Orijinalde numaralar bir "
  "sağa bir sola konumlanmış; pypdf hepsini metne alıyor, pdf-inspector temizliyor.")

d.add_heading("5.4 Sybil makalesi — pypdf önde", level=2)

P("Tıp dergisi düzeninde, her sayfada künye ve telif satırları bulunan, kutulu bilgi alanları "
  "içeren bir makale. Korpusun en zorlu belgesi çıktı ve pdf-inspector burada birden fazla "
  "yerde geriye düştü.")

P("En görünür sorun sayfa künyelerinin gövde metnine karışması. pdf-inspector on üç sayfada "
  "yirmi kez künye metnini gövde satırının içine soktu; bunların yedisi bir cümleyi ikiye "
  "böldü, biri kelimenin ortasına girdi. pypdf'de bu türden tek bir vaka yok.")

kod("pdf-inspector:\n"
    "  ...LDCTs that fit neither the true-positive nor true-\n"
    "  Copyright © 2026 American Society of Clinical Oncology. All rights reserved.\n"
    "  negative definition were excluded...")

P("Orijinal sayfada aynı cümle kesintisiz akıyor; araya giren satır PDF'in kendisinde o "
  "noktada bulunmuyor:")

gorsel("03_sybil_cumle.png",
       "Sybil makalesi, sayfa 4. Sol sütunda cümle kesintisiz: “…neither the true-positive "
       "nor true-negative definition were excluded…”.", 16.2)

P("İkinci sorun okuma sırasında. Üçüncü sayfada tek bir cümle üç parçaya bölünmüş, parçalar "
  "birbirinden uzağa dağılmış ve bir parçası başlık olarak işaretlenmiş. Dokuzuncu sayfada, "
  "yazar katkılarının listelendiği bölümde bir etiket tamamen kaybolmuş ve yazar adları yanlış "
  "katkı türlerine bağlanmış — yanlış atıf üretebilecek türde bir bozulma.")

P("Üçüncü sorun kaynakçada. pdf-inspector kaynak girdilerinin devam satırlarını, ait oldukları "
  "girdinin sekiz satır öncesine taşımış. pypdf'de kaynakça bozulmamış.")

P("Belgedeki kutulu bilgi alanı konusunda bir noktayı düzeltmek gerekiyor: kutuyu cümlenin "
  "ortasına yerleştirme hatası her iki kütüphanede de var. Aradaki fark kutunun içinde — pypdf "
  "kutu içindeki alt başlıkları ayrı satırlarda koruyor, pdf-inspector üçünü tek paragrafa "
  "eritiyor.")

P("pdf-inspector bu belgede de bağlı harflerde ve dipnot işaretlerinde kazandı, ama genel "
  "toplamda pypdf önde.")

# ================= 7. NEREDE HANGİSİ =================
d.add_heading("6. Hangi işte hangisi daha iyi", level=1)

P("Aşağıdaki tablo, dört belgede elle doğrulanan ve dokuz belgede otomatik ölçülen sonuçların "
  "birleşimidir. “Değişken” yazan satırlar, sonucun belgeden belgeye ciddi biçimde değiştiği "
  "durumları gösteriyor — bunlar entegrasyon kararı açısından en önemli satırlar.")

tablo(
    ["Konu", "pypdf", "pdf-inspector", "Sonuç"],
    [
        ["Hız", "0,13–1,19 sn", "0,02–0,13 sn", "pdf-inspector"],
        ["İçerik bütünlüğü", "Tam", "Tam", "Eşit"],
        ["Türkçe karakterler", "Korunuyor", "Korunuyor", "Eşit"],
        ["Bağlı harfler", "63–444 kalıntı", "0", "pdf-inspector"],
        ["Numaralı bölüm başlıkları", "Hiç üretmiyor", "6 belgenin 5'inde %100", "pdf-inspector"],
        ["Numarasız bölüm başlıkları", "Hiç üretmiyor", "12 başlıktan 1'i", "İkisi de zayıf"],
        ["Kaynakça filtresi", "Hiç çalışmıyor", "9 belgenin 7'sinde çalışıyor", "pdf-inspector, ama değişken"],
        ["Sütun okuma sırası", "Tüm belgelerde doğru", "2 belgede bozuk", "pypdf"],
        ["Sayfa künyesi", "Ayrı satırda tutuyor", "Bazen cümleye sokuyor", "pypdf"],
        ["Tablo yapısı", "Hiç kurmuyor", "Kuruyor ama %67'sinde hizasız", "Değişken"],
        ["Tekrarlanan sayfa damgası", "Her sayfada alıyor", "Bazen temizliyor", "pdf-inspector, ama tutarsız"],
        ["Çapraz referans boşluğu", "Doğru", "Sık düşürüyor", "pypdf"],
        ["Dipnot ayrımı", "Ayırmıyor", "Ayırmıyor", "İkisi de başarısız"],
        ["Satır sonu tirelemesi", "Çözmüyor", "Çözmüyor", "İkisi de başarısız"],
        ["Şekil içi yazılar", "Metne karışıyor", "Metne karışıyor", "İkisi de başarısız"],
        ["Taranmış belge sinyali", "Yok", "Sınıflandırma + OCR yönlendirme", "pdf-inspector"],
    ],
    [4.2, 3.6, 4.4, 3.2])

d.add_heading("6.1 Sonucun belgeye göre nasıl değiştiği", level=2)

P("Bu değerlendirmenin en önemli bulgusu, hangi kütüphanenin daha iyi olduğunun belgeye göre "
  "değişmesi. Aynı kütüphane bir belgede açık ara kazanırken diğerinde geriye düşebiliyor:")

tablo(
    ["Belge", "Genel sonuç", "Öne çıkan neden"],
    [
        ["BERT makalesi", "pdf-inspector önde", "Bağlı harf, tablo yapısı ve başlık tanıma birlikte kazandı"],
        ["VGG makalesi", "Karışık", "Damga temizliği kazandı, başlıklar tamamen kaybedildi"],
        ["Türkçe makale", "Başa baş", "Sayfa numaraları kazandı, sütun sırası ve tablolar kaybetti"],
        ["Sybil makalesi", "pypdf önde", "Künyenin cümlelere karışması, okuma sırası, kaynakça bozulması"],
    ],
    [3.6, 3.4, 8.4])

P("Sütun okuma sırası bozukluğu görülen iki belgenin ikisi de masaüstü yayıncılık yazılımıyla "
  "dizilmişti (InDesign ve Arbortext); elle incelenen LaTeX kaynaklı belgelerde bu kusura "
  "rastlanmadı. Bu bir eğilime işaret ediyor ama kesin bir kural olarak sunulamaz, çünkü elle "
  "incelenen belge sayısı dört ve bu ayrımın otomatik olarak ölçülmesi denendiğinde güvenilir "
  "sonuç alınamadı.")

P("Başlık tanıma ve tablo yapısı için ise böyle bir açıklama bulunamadı. Tablo kusuru dokuz "
  "belgenin dokuzunda da var. Başlık tarafında iki başarısızlık (VGG'nin alt başlıkları, "
  "Sybil'in büyük harfli başlıkları) belgelerin dizgi yazılımıyla açıklanamıyor; Sybil'de "
  "tanınan ve kaçırılan başlıklar aynı puntoda dizilmiş olduğu için tipografiyle de "
  "açıklanamıyor.")

# ================= 8. İDDİALAR =================
d.add_heading("7. pdf-inspector'ın iddiaları", level=1)

P("Kütüphanenin belgelerinde öne sürülen özellikler, test ettiğimiz dosyalar üzerinde "
  "aşağıdaki gibi karşılık buldu.")

tablo(
    ["İddia", "Durum", "Not"],
    [
        ["Belge türü sınıflandırma ve OCR yönlendirme", "Karşılıyor",
         "Taranmış belgeyi 0,95 güvenle doğru sınıflandırdı"],
        ["Metin tabanlı PDF'lerde 200 ms altı süre", "Aşıyor",
         "Ölçülen aralık 21–132 ms"],
        ["Konum farkındalıklı çıkarma", "Karşılıyor",
         "Konum verisi API üzerinden erişilebilir"],
        ["Çok sütunlu okuma sırası", "Kısmen",
         "İki belgede sütunlar satır satır karıştı"],
        ["Markdown dönüşümü (başlık, liste, biçim)", "Kısmen",
         "Numaralı başlıklarda güçlü, numarasız başlıklarda zayıf"],
        ["Tablo tespiti", "Kısmen",
         "Tablolar üretiliyor, %67'sinde sütun hizası bozuk"],
        ["Bağlı harf ve kodlama düzeltmesi", "Karşılıyor",
         "Tüm belgelerde sıfır kalıntı"],
        ["Dipnot işleme", "Karşılamıyor",
         "Belge dipnotları gövdeden ayrılmıyor"],
        ["Bozuk kodlama tespiti", "Test edilmedi",
         "Korpusta uygun belge yoktu"],
    ],
    [5.2, 2.6, 7.6])

P("Genel değerlendirme: kütüphane hız, sınıflandırma ve tipografik temizlik iddialarını "
  "fazlasıyla karşılıyor. Yapı çıkarma iddialarını (başlık, tablo, okuma sırası) ise kısmen "
  "karşılıyor — özellik çalışıyor, ama güvenilirliği belgeden belgeye değişiyor.")

# ================= 10. SINIRLAR =================
d.add_heading("8. Bu çalışmanın sınırları", level=1)

P("Sonuçların nasıl okunması gerektiği açısından aşağıdaki noktalar belirtilmelidir.")

madde("Elle inceleme dört belgeyle sınırlı. Diğer beş belge yalnızca otomatik ölçümlere dahil.")
madde("Sütun okuma sırası bozukluğunu otomatik tespit etmek için üç ayrı yöntem denendi, "
      "hiçbiri güvenilir sonuç vermedi. Bu kusur yalnızca gözle doğrulandığı yerlerde "
      "raporlandı; diğer belgelerde bulunmadığı değil, aranmadığı anlaşılmalıdır.")
madde("Dizgi yazılımı ile okuma sırası arasındaki bağlantı dört belgelik bir gözleme dayanıyor; "
      "eğilim olarak sunuldu, kural olarak değil.")
madde("Bozuk kodlamalı ve şifreli PDF davranışı test edilmedi; korpusta bu türden belge yoktu.")
madde("Görsel işleme her iki kütüphanenin de kapsamı dışında. Platformda şekil analizi ayrı bir "
      "hat üzerinden yürüdüğü için bu değişiklikten etkilenmez.")

# ================= 11. SONUÇ =================
d.add_heading("9. Sonuç", level=1)

P("pdf-inspector iyi yazılmış, hızlı ve belirli işlerde pypdf'ten açıkça üstün bir kütüphane. "
  "Bağlı harf temizliği kusursuz, hız farkı yadsınamaz, belge sınıflandırma özelliği "
  "pypdf'te hiç bulunmayan bir yetenek. Bölüm başlıklarını çıkarabilmesi ise sistemde şu anda "
  "ölü duran bir doğrulama mekanizmasını canlandırma potansiyeli taşıyor.")

P("Buna karşılık yapı çıkarma tarafındaki davranışı her belgede aynı kalitede değil. Numaralı "
  "bölüm başlıklarını altı belgenin beşinde kusursuz buluyor; buna karşılık VGG'de on bir alt "
  "başlığın hepsini, Sybil'de on iki başlıktan on birini kaçırdı. Ürettiği tabloların üçte "
  "ikisinde sütun hizası bozuk, iki belgede sütun okuma sırasını karıştırdı, bir belgede sayfa "
  "künyelerini cümlelerin ortasına soktu. Bu kusurların ortak özelliği, çıktıya bakarak fark "
  "edilememeleri.")

P("Kararı zorlaştıran asıl mesele kusurların türü. pypdf'in ürettiği bozukluklar bakınca "
  "anlaşılıyor: tablo dağılmış, başlık yok, sayılar düz metin hâlinde. pdf-inspector'ın "
  "ürettiği bozukluklar ise düzgün görünüyor — geçerli bir Markdown tablosu, geçerli bir "
  "başlık hiyerarşisi. Yanlış oldukları ancak orijinal PDF ile karşılaştırılınca anlaşılıyor. "
  "Otomatik bir işlem hattında ikinci tür hata daha tehlikelidir.")

P("Bu nedenle bu aşamada pypdf'i çıkarıp yerine pdf-inspector koymak net bir iyileşme "
  "getirmiyor; kazanılacaklar ile kaybedilecekler aynı büyüklükte. Öte yandan kütüphaneyi "
  "tamamen elemek de doğru olmaz: şu anda kaynakça filtresi PDF kaynakları için hiç "
  "çalışmıyor ve pdf-inspector bunu dokuz belgenin yedisinde çalışır hâle getiriyor. "
  "Gövde metnini pypdf'ten, başlık bilgisini pdf-inspector'dan alan karma bir kullanım, iki "
  "kütüphanenin güçlü yanlarını riske girmeden birleştirebilir. Hangi yol seçilirse seçilsin, "
  "karar öncesinde daha geniş bir belge kümesiyle doğrulama yapılması yerinde olur.")

d.add_paragraph()
P("Ham ölçüm verileri sonuclar.csv dosyasında, elle inceleme notları bulgular.md dosyasında, "
  "iki kütüphanenin ham çıktıları out/ klasöründe bulunmaktadır.",
  italik=True, boyut=9.5)

adaylar = ["RAPOR.docx"] + [f"RAPOR_v{i}.docx" for i in range(2, 30)]
for ad in adaylar:
    try:
        d.save(ad)
        print(f"{ad} olusturuldu")
        if ad != "RAPOR.docx":
            print("(RAPOR.docx Word'de acik oldugu icin uzerine yazilamadi)")
        break
    except PermissionError:
        continue
else:
    print("Hicbir dosya adina yazilamadi. Word'de acik olan RAPOR dosyalarini kapat.")
